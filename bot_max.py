"""
Бот дежурств — версия для мессенджера MAX (перенос с python-telegram-bot на maxapi).

Бизнес-логика (DutyScheduleGenerator, круг дежурств, расписание уведомлений)
перенесена без изменений. Изменился только транспортный слой: вместо
python-telegram-bot используется библиотека maxapi (https://github.com/love-apples/maxapi).

Запуск:
    MAX_BOT_TOKEN=ваш_токен python bot_max.py
(токен читается из переменной окружения — см. .env.example)
"""

import asyncio
import contextlib
import json
import logging
import os
import shutil
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional

# Опционально: загрузка .env, если установлен python-dotenv
with contextlib.suppress(ImportError):
    from dotenv import load_dotenv

    load_dotenv()

import pytz
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger
from docx import Document

from maxapi import Bot, Dispatcher, F
from maxapi.context.base import BaseContext
from maxapi.context.state_machine import State, StatesGroup
from maxapi.enums.parse_mode import TextFormat
from maxapi.filters.command import Command, CommandStart
from maxapi.types.attachments.buttons.callback_button import CallbackButton
from maxapi.types.input_media import InputMedia
from maxapi.types.updates.bot_started import BotStarted
from maxapi.types.updates.message_callback import MessageCallback
from maxapi.types.updates.message_created import MessageCreated
from maxapi.utils.inline_keyboard import InlineKeyboardBuilder

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Отдельный лог для аудита действий (ТЗ п.3, "Логирование": кто/когда/что сделал)
audit_logger = logging.getLogger("audit")
audit_logger.setLevel(logging.INFO)
audit_logger.propagate = False
_audit_handler = logging.FileHandler("audit.log", encoding="utf-8")
_audit_handler.setFormatter(logging.Formatter('%(asctime)s | %(message)s'))
audit_logger.addHandler(_audit_handler)


def audit(user_id: str, username: Optional[str], action: str, details: str = ""):
    """Пишет строку в audit.log: кто, когда, что сделал (ТЗ п.3)."""
    who = f"{user_id}(@{username})" if username else user_id
    audit_logger.info(f"{who} | {action}" + (f" | {details}" if details else ""))

# Часовой пояс (Москва)
MOSCOW_TZ = pytz.timezone('Europe/Moscow')

# Персональные данные (ФИО, телефоны, логин/пароль админки) НЕ хранятся в коде —
# они читаются либо из config.json (локально/на VPS), либо из переменной
# окружения CONFIG_JSON (на хостингах вроде Railway/Render/Fly, где нет
# доступа к файловой системе репозитория, только env-переменные).
# Шаблон структуры: config.example.json.
CONFIG_PATH = Path(__file__).resolve().parent / "config.json"


def _load_config() -> dict:
    env_config = os.environ.get("CONFIG_JSON")
    if env_config:
        return json.loads(env_config)

    if not CONFIG_PATH.exists():
        raise FileNotFoundError(
            f"Не найден {CONFIG_PATH} и не задана переменная окружения CONFIG_JSON. "
            "Локально: скопируйте config.example.json в config.json и заполните реальными "
            "данными. На хостинге без файловой системы: положите содержимое config.json "
            "целиком (в одну строку) в переменную окружения CONFIG_JSON."
        )
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


_config = _load_config()

# Админ-аккаунт
ADMIN_CREDENTIALS = {
    "login": _config["admin_login"],
    "password": _config["admin_password"]
}

# Соответствие username в MAX сотрудникам
USERNAME_TO_EMPLOYEE = _config["username_to_employee"]

# Телефоны сотрудников
EMPLOYEE_PHONES = _config["employee_phones"]

# Строгая последовательность дежурств по кругу
DUTY_ROTATION_CIRCLE = _config["duty_rotation_circle"]

# Исходный базовый список для хранения ручных правок админов
DUTY_SCHEDULE = []

# user_id администраторов, которым при первом запуске бота автоматически
# выдаются права админа (ТЗ п.3, "Безопасность": проверка по списку ADMIN_IDS).
# Основной способ входа в админку — по-прежнему /admin логин пароль;
# ADMIN_IDS лишь заранее "пред-авторизует" перечисленных пользователей.
ADMIN_IDS = [int(x) for x in _config.get("admin_ids", [])]

# Защита от посторонних (ТЗ п.2.4). По умолчанию выключена, чтобы не
# заблокировать случайно всю команду при первом деплое — включается явно.
WHITELIST_MODE = bool(_config.get("whitelist_mode", False))

# Настройки опроса по итогам дежурства (ТЗ п.2.2)
SURVEY_CONFIG = _config.get("survey", {
    "send_hour": 8,
    "send_minute": 0,
    "questions": {
        "quality": {"text": "Как прошло дежурство?", "options": ["Отлично", "Хорошо", "Были сложности", "Плохо"]},
        "incidents": {"text": "Были ли инциденты?", "options": ["Да", "Нет"]},
        "zgd": {"text": "Были ли ЗГД?", "options": ["Да", "Нет"]},
    }
})

# Маска имени файла автопротокола (ТЗ п.2.5)
PROTOCOL_FILENAME_MASK = _config.get(
    "protocol_filename_mask", "Протокол Разногласий_{date}_Смена{shift_number}.docx"
)

# Обновляется вручную при каждом релизе — по /time можно однозначно проверить,
# какая версия кода реально работает на хостинге (без гадания по редеплою).
BOT_CODE_VERSION = "2026-07-10-clear-stale-keyboards"


class DutyScheduleGenerator:
    """Генератор графика дежурств с поддержкой полного цикла из 12 недель.

    Не зависит от мессенджера — перенесено без изменений.
    """

    def __init__(self, schedule_data: List[Dict]):
        self.schedule_data = schedule_data
        self.schedule = {}
        self.initialize_schedule()

    def initialize_schedule(self):
        for duty in self.schedule_data:
            self.schedule[duty["date"]] = {
                "employees": duty["employees"],
                "phones": duty["phones"],
                "is_pair": duty["is_pair"],
                "date_obj": duty["date_obj"]
            }
        logger.info(f"Загружен график на {len(self.schedule)} недель")
        self.remove_past_duties()

    def remove_past_duties(self):
        today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
        dates_to_remove = []

        for date_str, duty in self.schedule.items():
            if duty["date_obj"] < today:
                dates_to_remove.append(date_str)

        for date_str in dates_to_remove:
            del self.schedule[date_str]
            self.schedule_data = [d for d in self.schedule_data if d["date"] != date_str]

        if dates_to_remove:
            logger.info(f"Удалено {len(dates_to_remove)} прошедших ручных дежурств")

    def _get_upcoming_saturdays(self, count: int = 12) -> List[datetime]:
        saturdays = []
        today = datetime.now(MOSCOW_TZ).replace(hour=0, minute=0, second=0, microsecond=0).replace(tzinfo=None)

        days_ahead = 5 - today.weekday()
        if days_ahead < 0:
            days_ahead += 7

        current_saturday = today + timedelta(days=days_ahead)

        for _ in range(count):
            saturdays.append(current_saturday)
            current_saturday += timedelta(days=7)

        return saturdays

    def _generate_dynamic_schedule(self) -> Dict[str, Dict]:
        base_date = datetime(2026, 5, 30)
        base_index = 0

        now_moscow = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
        all_saturdays = self._get_upcoming_saturdays(count=13)

        if now_moscow.weekday() == 5 and now_moscow.hour >= 8:
            active_saturdays = all_saturdays[1:13]
        else:
            active_saturdays = all_saturdays[:12]

        dynamic_schedule = {}

        for sat in active_saturdays:
            date_str = sat.strftime("%d.%m.%Yг.")

            if date_str in self.schedule:
                dynamic_schedule[date_str] = self.schedule[date_str]
            else:
                weeks_diff = int((sat - base_date).days / 7)
                employee_index = (base_index + weeks_diff) % len(DUTY_ROTATION_CIRCLE)
                employee_name = DUTY_ROTATION_CIRCLE[employee_index]
                phone = EMPLOYEE_PHONES.get(employee_name, "не указан")

                dynamic_schedule[date_str] = {
                    "employees": [employee_name],
                    "phones": [phone],
                    "is_pair": False,
                    "date_obj": sat
                }

        return dynamic_schedule

    def get_schedule_text(self) -> str:
        text = "📅 <b>АКТУАЛЬНЫЙ ГРАФИК ДЕЖУРСТВ</b>\n\n"

        current_schedule = self._generate_dynamic_schedule()
        duties_list = sorted(current_schedule.items(), key=lambda x: x[1]["date_obj"])

        if not duties_list:
            text += "Нет запланированных дежурств\n"
        else:
            for i, (date_str, duty) in enumerate(duties_list):
                if i == 0:
                    text += f"<b>{date_str} (Ближайшее)</b>\n"
                else:
                    text += f"{date_str}\n"

                if duty["is_pair"]:
                    text += f"{duty['employees'][0]} + {duty['employees'][1]}\n"
                    text += f"{duty['phones'][0]} + {duty['phones'][1]}\n\n"
                else:
                    text += f"{duty['employees'][0]}\n"
                    text += f"{duty['phones'][0]}\n\n"

        text += f"<i>Актуально на: {datetime.now(MOSCOW_TZ).strftime('%d.%m.%Y %H:%M')}</i>"
        return text

    def get_employee_schedule(self, employee_name: str) -> List[Dict]:
        result = []
        current_schedule = self._generate_dynamic_schedule()

        for date_str, duty in current_schedule.items():
            if employee_name in duty["employees"]:
                result.append({
                    "date": date_str,
                    "employees": duty["employees"],
                    "phones": duty["phones"],
                    "is_pair": duty["is_pair"],
                    "date_obj": duty["date_obj"]
                })
        return sorted(result, key=lambda x: x["date_obj"])

    def get_next_duty(self, employee_name: str) -> Optional[Dict]:
        duties = self.get_employee_schedule(employee_name)
        return duties[0] if duties else None

    def get_todays_duty(self) -> Optional[Dict]:
        now_moscow = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
        today_str = now_moscow.strftime("%d.%m.%Yг.")

        if today_str in self.schedule:
            return self.schedule[today_str]

        base_date = datetime(2026, 5, 30)
        all_saturdays = self._get_upcoming_saturdays(count=5)
        for sat in all_saturdays:
            if sat.strftime("%d.%m.%Yг.") == today_str:
                weeks_diff = int((sat - base_date).days / 7)
                employee_index = weeks_diff % len(DUTY_ROTATION_CIRCLE)
                employee_name = DUTY_ROTATION_CIRCLE[employee_index]
                return {
                    "employees": [employee_name],
                    "phones": [EMPLOYEE_PHONES.get(employee_name)],
                    "is_pair": False,
                    "date_obj": sat
                }
        return None

    def add_duty(self, date_str: str, employees: List[str], phones: List[str], is_pair: bool):
        try:
            date_str_clean = date_str.replace("г.", "").strip()
            date_obj = datetime.strptime(date_str_clean, "%d.%m.%Y")
            today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)

            if date_obj < today and date_str != today.strftime("%d.%m.%Yг."):
                return False, "Дата должна быть в будущем или текущей"

            if not date_str.endswith("г."):
                date_str += "г."

            self.schedule[date_str] = {
                "employees": employees,
                "phones": phones,
                "is_pair": is_pair,
                "date_obj": date_obj
            }

            self.schedule_data.append({
                "date": date_str,
                "date_obj": date_obj,
                "employees": employees,
                "phones": phones,
                "is_pair": is_pair
            })

            logger.info(f"Добавлено ручное дежурство: {date_str} - {employees}")
            return True, "Дежурство успешно добавлено"
        except Exception as e:
            logger.error(f"Ошибка добавления дежурства: {e}")
            return False, f"Ошибка: {str(e)}"

    def remove_duty(self, date_str: str) -> bool:
        if not date_str.endswith("г."):
            date_str += "г."
        if date_str in self.schedule:
            del self.schedule[date_str]
            self.schedule_data = [d for d in self.schedule_data if d["date"] != date_str]
            logger.info(f"Удалено ручное дежурство: {date_str}")
            return True
        return False

    def update_employee_phone(self, employee_name: str, new_phone: str) -> bool:
        global EMPLOYEE_PHONES
        if employee_name in EMPLOYEE_PHONES:
            EMPLOYEE_PHONES[employee_name] = new_phone
            logger.info(f"Обновлен телефон {employee_name}: {new_phone}")
            return True
        return False

    def add_employee(self, employee_name: str, phone: str) -> bool:
        global EMPLOYEE_PHONES
        if employee_name not in EMPLOYEE_PHONES:
            EMPLOYEE_PHONES[employee_name] = phone
            logger.info(f"Добавлен сотрудник: {employee_name} - {phone}")
            return True
        return False

    def remove_employee(self, employee_name: str) -> bool:
        global EMPLOYEE_PHONES
        if employee_name in EMPLOYEE_PHONES:
            del EMPLOYEE_PHONES[employee_name]
            logger.info(f"Удален сотрудник: {employee_name}")
            return True
        return False


class AdminWizard(StatesGroup):
    """Состояния пошагового мастера админ-панели (замена context.user_data из PTB)."""
    wait_date = State()
    wait_phones = State()
    awaiting_duty_remove = State()
    awaiting_employee_add = State()
    awaiting_employee_remove = State()
    awaiting_phone_edit = State()
    awaiting_protocol_upload = State()
    awaiting_protocol_pin = State()


class SurveyWizard(StatesGroup):
    """Состояние опроса по итогам дежурства (ТЗ п.2.2) — только текстовый
    шаг «Замечания», остальные вопросы отвечаются кнопками через callback."""
    awaiting_remarks = State()


class DutyBot:
    def __init__(self, token: Optional[str] = None):
        # Bot() без аргумента сам читает токен из переменной окружения MAX_BOT_TOKEN
        self.bot = Bot(token=token) if token else Bot()
        self.dp = Dispatcher()

        self.schedule_generator = DutyScheduleGenerator(DUTY_SCHEDULE)
        self.user_data_file = "user_data.json"
        self.shifts_file = "shifts_history.json"
        self.protocol_file_path = "Протокол разногласий — пример.docx"
        self.protocol_pinned_message_id = None
        self.download_dir = Path("downloads")
        self.protocol_dir = Path("protocols")
        self.scheduler = None
        self.load_user_data()
        self.load_shifts()
        self._seed_admin_ids()
        self._register_handlers()

    def _seed_admin_ids(self):
        """Пользователи из ADMIN_IDS (config.json) получают права админа
        сразу, без /admin login (ТЗ п.3, «Безопасность»)."""
        changed = False
        for admin_id in ADMIN_IDS:
            uid = str(admin_id)
            if uid not in self.user_data:
                self.user_data[uid] = {
                    "username": None, "first_name": None, "last_name": None,
                    "display_name": "Админ (ADMIN_IDS)", "notifications": True,
                    "selected_employee": None, "registered_at": datetime.now().isoformat(),
                    "last_active": datetime.now().isoformat(), "is_admin": True,
                    "whitelisted": True,
                }
                changed = True
            elif not self.user_data[uid].get("is_admin"):
                self.user_data[uid]["is_admin"] = True
                changed = True
        if changed:
            self.save_user_data()

    # ================= РЕГИСТРАЦИЯ ХЕНДЛЕРОВ =================

    def _register_handlers(self):
        dp = self.dp

        dp.bot_started()(self.on_bot_started)
        dp.message_created(CommandStart())(self.start)
        dp.message_created(Command("admin"))(self.admin_login)
        dp.message_created(Command("test_wednesday"))(self.send_test_wednesday)
        dp.message_created(Command("test_friday"))(self.send_test_friday)
        dp.message_created(Command("test_saturday"))(self.send_test_saturday)
        dp.message_created(Command("test_survey"))(self.send_test_survey)
        dp.message_created(Command("test_user"))(self.test_notification_for_user)
        dp.message_created(Command("users"))(self.check_users_status)
        dp.message_created(Command("enable_all"))(self.enable_notifications_all)
        dp.message_created(Command("test_send"))(self.test_send_to_user)
        dp.message_created(Command("time"))(self.check_time)
        dp.message_created(Command("fix"))(self.fix_all_users)
        dp.message_created(Command("stats"))(self.cmd_stats)
        dp.message_created(Command("contact"))(self.cmd_contact)
        dp.message_created(Command("set_phone"))(self.cmd_set_phone)
        dp.message_created(Command("set_schedule"))(self.cmd_set_schedule)
        dp.message_created(Command("myid"))(self.cmd_myid)

        dp.message_callback()(self.on_callback)

        # Пошаговый мастер добавления дежурства и правки сотрудников
        dp.message_created(AdminWizard.wait_date)(self.wizard_wait_date)
        dp.message_created(AdminWizard.wait_phones)(self.wizard_wait_phones)
        dp.message_created(AdminWizard.awaiting_duty_remove)(self.wizard_duty_remove)
        dp.message_created(AdminWizard.awaiting_employee_add)(self.wizard_employee_add)
        dp.message_created(AdminWizard.awaiting_employee_remove)(self.wizard_employee_remove)
        dp.message_created(AdminWizard.awaiting_phone_edit)(self.wizard_phone_edit)
        dp.message_created(AdminWizard.awaiting_protocol_upload)(self.wizard_protocol_upload)
        dp.message_created(AdminWizard.awaiting_protocol_pin)(self.wizard_protocol_pin)
        dp.message_created(SurveyWizard.awaiting_remarks)(self.wizard_survey_remarks)

        # Файл прислан без активного состояния мастера — подскажем, что делать
        # (должно идти раньше общего текстового фолбэка)
        dp.message_created(F.message.body.attachments)(self.on_document)

        # Фолбэк для любого текста вне активного состояния мастера
        dp.message_created(None)(self.on_plain_text)

    # ================= ПЛАНИРОВЩИК И УВЕДОМЛЕНИЯ =================

    async def _start_scheduler(self):
        self.scheduler = AsyncIOScheduler(timezone=MOSCOW_TZ)

        self.scheduler.add_job(
            self.send_wednesday_notification,
            CronTrigger(day_of_week='wed', hour=18, minute=0, second=0, timezone=MOSCOW_TZ),
            id='wednesday_notification',
            replace_existing=True
        )
        self.scheduler.add_job(
            self.send_friday_notification_all,
            CronTrigger(day_of_week='fri', hour=18, minute=0, second=0, timezone=MOSCOW_TZ),
            id='friday_notification_all',
            replace_existing=True
        )
        self.scheduler.add_job(
            self.send_saturday_notification_all,
            CronTrigger(day_of_week='sat', hour=10, minute=0, second=0, timezone=MOSCOW_TZ),
            id='saturday_notification_all',
            replace_existing=True
        )
        self.scheduler.add_job(
            self.send_shift_survey,
            CronTrigger(day_of_week='sat', hour=SURVEY_CONFIG.get("send_hour", 8),
                        minute=SURVEY_CONFIG.get("send_minute", 0), second=0, timezone=MOSCOW_TZ),
            id='shift_survey',
            replace_existing=True
        )

        self.scheduler.start()
        logger.info(
            f"Планировщик задач запущен: среда 18:00, пятница 18:00, "
            f"суббота {SURVEY_CONFIG.get('send_hour', 8):02d}:{SURVEY_CONFIG.get('send_minute', 0):02d} (опрос), суббота 10:00"
        )

    async def send_wednesday_notification(self):
        try:
            today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
            if today.weekday() != 2:
                logger.warning(f"send_wednesday_notification вызван не в среду! День недели: {today.weekday()}")
                return

            logger.info(f"Запуск send_wednesday_notification в среду {today.strftime('%d.%m.%Y %H:%M')}")
            saturday = today + timedelta(days=3)

            current_schedule = self.schedule_generator._generate_dynamic_schedule()
            duty_saturday = next((d for d in current_schedule.values() if d["date_obj"].date() == saturday.date()),
                                 None)

            if not duty_saturday:
                message = (
                    f"🔔 <b>НАПОМИНАНИЕ О ДЕЖУРСТВЕ В СУББОТУ</b>\n\n"
                    f"📅 <b>{saturday.strftime('%d.%m.%Y')}</b>\n\n"
                    f"⚠️ <b>В эту субботу дежурных нет</b>\n\n"
                    f"✅ Можно отдыхать!\n\n"
                    f"<i>Следующее напоминание: пятница в 18:00</i>"
                )
            else:
                if duty_saturday["is_pair"]:
                    duty_text = f"{duty_saturday['employees'][0]} + {duty_saturday['employees'][1]}"
                    phones_text = f"{duty_saturday['phones'][0]} + {duty_saturday['phones'][1]}"
                else:
                    duty_text = f"{duty_saturday['employees'][0]}"
                    phones_text = f"{duty_saturday['phones'][0]}"

                message = (
                    f"🔔 <b>НАПОМИНАНИЕ О ДЕЖУРСТВЕ В СУББОТУ</b>\n\n"
                    f"📅 <b>Дата:</b> {saturday.strftime('%d.%m.%Y')}\n"
                    f"👤 <b>Дежурит:</b> {duty_text}\n"
                    f"📞 <b>Телефоны:</b> {phones_text}\n\n"
                    f"⏰ <b>Время:</b> 6:50 - 8:00\n"
                    f"📍 <b>Место:</b> кабинет 6002, 6 этаж, АДЦ\n\n"
                    f"📋 <b>Инструкция:</b>\n"
                    f"• В пятницу до 17:00 позвонить в приемную: 5600\n"
                    f"• Прийти в субботу к 6:50 в АДЦ\n"
                    f"• Взять ключ на охране от кубов\n"
                    f"• Сфотографировать открытый кабинет\n"
                    f"• Находиться там до 8:00\n\n"
                    f"<i>Следующее напоминание: пятница в 18:00</i>"
                )

            await self._send_notification_to_all_users(message, "среда")
        except Exception as e:
            logger.error(f"Ошибка отправки уведомления в среду: {e}")

    async def send_friday_notification_all(self):
        try:
            today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
            if today.weekday() != 4:
                logger.warning(f"send_friday_notification_all вызван не в пятницу! День недели: {today.weekday()}")
                return

            logger.info(f"Запуск send_friday_notification_all в пятницу {today.strftime('%d.%m.%Y %H:%M')}")
            tomorrow = today + timedelta(days=1)

            current_schedule = self.schedule_generator._generate_dynamic_schedule()
            duty_tomorrow = next((d for d in current_schedule.values() if d["date_obj"].date() == tomorrow.date()),
                                 None)

            if not duty_tomorrow:
                message = (
                    f"🔔 <b>НАПОМИНАНИЕ О ЗАВТРАШНЕМ ДЕЖУРСТВЕ</b>\n\n"
                    f"📅 <b>Завтра ({tomorrow.strftime('%d.%m.%Y')}) дежурных нет</b>\n\n"
                    f"✅ Можете не беспокоиться!\n\n"
                    f"<i>Следующее напоминание: суббота в 10:00</i>"
                )
            else:
                if duty_tomorrow["is_pair"]:
                    duty_text = f"{duty_tomorrow['employees'][0]} + {duty_tomorrow['employees'][1]}"
                    phones_text = f"{duty_tomorrow['phones'][0]} + {duty_tomorrow['phones'][1]}"
                else:
                    duty_text = f"{duty_tomorrow['employees'][0]}"
                    phones_text = f"{duty_tomorrow['phones'][0]}"

                message = (
                    f"🔔 <b>НАПОМИНАНИЕ О ЗАВТРАШНЕМ ДЕЖУРСТВЕ</b>\n\n"
                    f"📅 <b>Завтра ({tomorrow.strftime('%d.%m.%Y')}) дежурит:</b>\n"
                    f"👤 {duty_text}\n"
                    f"📞 {phones_text}\n\n"
                    f"⏰ <b>Время:</b> 6:50 - 8:00\n"
                    f"📍 <b>Место:</b> кабинет 6002, 6 этаж, АДЦ\n\n"
                    f"⚠️ <b>ВАЖНО! СЕГОДНЯ ДО 19:00:</b>\n"
                    f"• Дежурным позвонить в приемную: 5600\n"
                    f"• Сообщить о дежурстве\n"
                    f"• Попросить оставить ключи на вахте\n\n"
                    f"📋 <b>План на завтра:</b>\n"
                    f"• Прийти в АДЦ к 6:50\n"
                    f"• Взять ключ на охране от кубов\n"
                    f"• Открыть кабинет 6002\n"
                    f"• Сфотографировать открытый кабинет\n"
                    f"• Находиться там до 8:00\n"
                    f"• Оформить протокол разногласий\n\n"
                    f"<i>Следующее напоминание: суббота в 10:00</i>"
                )

            await self._send_notification_to_all_users(message, "пятница")
        except Exception as e:
            logger.error(f"Ошибка отправки уведомления в пятницу: {e}")

    async def send_saturday_notification_all(self):
        try:
            today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
            if today.weekday() != 5:
                logger.warning(f"send_saturday_notification_all вызван не в субботу! День недели: {today.weekday()}")
                return

            logger.info(f"Запуск send_saturday_notification_all в субботу {today.strftime('%d.%m.%Y %H:%M')}")
            duty_today = self.schedule_generator.get_todays_duty()

            if not duty_today:
                message = (
                    f"🔔 <b>ИНФОРМАЦИЯ О ДЕЖУРСТВЕ</b>\n\n"
                    f"📅 <b>Сегодня ({today.strftime('%d.%m.%Y')}) дежурных нет</b>\n\n"
                    f"✅ Всем хороших выходных!\n\n"
                    f"<i>Следующее напоминание: среда в 18:00</i>"
                )
            else:
                if duty_today["is_pair"]:
                    duty_text = f"{duty_today['employees'][0]} + {duty_today['employees'][1]}"
                    phones_text = f"{duty_today['phones'][0]} + {duty_today['phones'][1]}"
                else:
                    duty_text = f"{duty_today['employees'][0]}"
                    phones_text = f"{duty_today['phones'][0]}"

                message = (
                    f"🔔 <b>ИТОГИ ДЕЖУРСТВА</b>\n\n"
                    f"📅 <b>Сегодня ({today.strftime('%d.%m.%Y')}) дежурили:</b>\n"
                    f"👤 {duty_text}\n"
                    f"📞 {phones_text}\n\n"
                    f"✅ <b>Дежурство завершилось в 8:00</b>\n\n"
                    f"📋 <b>Напоминание дежурным:</b>\n"
                    f"• Не забудьте оформить протокол разногласий\n"
                    f"• Протокол оставить у Е.С. Денисовой\n\n"
                    f"<i>Следующее напоминание: среда в 18:00</i>"
                )

            await self._send_notification_to_all_users(message, "суббота")
        except Exception as e:
            logger.error(f"Ошибка отправки уведомления в субботу: {e}")

    def _get_user_ids_for_employee(self, employee_name: str) -> List[str]:
        return [uid for uid, info in self.user_data.items() if info.get("selected_employee") == employee_name]

    def _get_shift(self, shift_number: int) -> Optional[Dict]:
        return next((s for s in self.shifts if s.get("shift_number") == shift_number), None)

    async def send_shift_survey(self, force: bool = False) -> Dict:
        """Опрос по итогам дежурства (ТЗ п.2.2) — отправляется дежурному(ым) в
        SURVEY_CONFIG['send_hour']:SURVEY_CONFIG['send_minute'] по субботам.

        force=True (команда /test_survey) снимает проверку дня недели и
        повторной отправки — удобно для ручного тестирования в любой день.

        Возвращает словарь {ok, sent, employees, reason} — используется
        /test_survey, чтобы явно показать, ушёл ли опрос кому-то реально,
        а не просто молча отчитаться об «успехе»."""
        try:
            today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
            if today.weekday() != 5 and not force:
                logger.warning(f"send_shift_survey вызван не в субботу! День недели: {today.weekday()}")
                return {"ok": False, "sent": 0, "employees": [], "reason": "не суббота"}

            duty_today = self.schedule_generator.get_todays_duty()
            if not duty_today and force:
                # Тестовый режим в будний день: берём ближайшее дежурство по кругу
                current_schedule = self.schedule_generator._generate_dynamic_schedule()
                if current_schedule:
                    duty_today = min(current_schedule.values(), key=lambda d: d["date_obj"])
            if not duty_today:
                logger.info("Опрос по дежурству: дежурных нет, опрос не отправляется")
                return {"ok": False, "sent": 0, "employees": [], "reason": "нет дежурных в графике"}

            date_str = today.strftime("%d.%m.%Yг.")
            if any(s["date"] == date_str for s in self.shifts) and not force:
                logger.info(f"Опрос по смене {date_str} уже отправлялся, повторно не отправляем")
                return {"ok": False, "sent": 0, "employees": [], "reason": "опрос по этой смене уже отправлялся"}

            shift_number = len(self.shifts) + 1
            shift = {
                "shift_number": shift_number,
                "date": date_str,
                "employees": duty_today["employees"],
                "survey": {},
                "protocol_file": None,
                "completed": False,
            }
            self.shifts.append(shift)
            self.save_shifts()

            question = SURVEY_CONFIG["questions"]["quality"]
            kb = InlineKeyboardBuilder()
            for i, option in enumerate(question["options"]):
                kb.row(CallbackButton(text=option, payload=f"survey|{shift_number}|quality|{i}"))

            sent_to = 0
            unlinked_employees = []
            for emp in duty_today["employees"]:
                recipients = self._get_user_ids_for_employee(emp)
                if not recipients:
                    unlinked_employees.append(emp)
                for uid in recipients:
                    try:
                        await self.bot.send_message(
                            user_id=int(uid),
                            text=f"🔔 <b>ОПРОС ПО ИТОГАМ ДЕЖУРСТВА</b>\n\n{question['text']}",
                            attachments=[kb.as_markup()],
                            format=TextFormat.HTML
                        )
                        sent_to += 1
                    except Exception as e:
                        logger.error(f"Не удалось отправить опрос пользователю {uid}: {e}")

            if unlinked_employees:
                logger.warning(
                    f"Опрос по смене №{shift_number}: сотрудники без привязанного "
                    f"пользователя бота — {', '.join(unlinked_employees)}"
                )

            logger.info(f"Опрос по смене №{shift_number} ({date_str}) отправлен {sent_to} получателям")
            audit("system", None, "survey_sent", f"смена №{shift_number}, {date_str}, получателей: {sent_to}")
            return {
                "ok": True, "sent": sent_to, "employees": duty_today["employees"],
                "unlinked": unlinked_employees, "shift_number": shift_number,
            }
        except Exception as e:
            logger.error(f"Ошибка отправки опроса по дежурству: {e}")
            return {"ok": False, "sent": 0, "employees": [], "reason": str(e)}

    async def handle_survey_callback(self, event: MessageCallback, context: BaseContext):
        """Обработка нажатий кнопок опроса: payload вида survey|{смена}|{вопрос}|{индекс}.
        event.answer() уже вызван в on_callback перед диспетчеризацией."""
        payload = event.callback.payload

        try:
            _, shift_number_str, question_key, option_index_str = payload.split("|")
            shift_number = int(shift_number_str)
            option_index = int(option_index_str)
        except (ValueError, AttributeError) as e:
            logger.error(f"Некорректный payload опроса '{payload}': {e}")
            if event.message is not None:
                await event.message.answer("❌ Не удалось обработать кнопку опроса.", format=TextFormat.HTML)
            return

        if event.message is None:
            return

        try:
            shift = self._get_shift(shift_number)
            if not shift:
                logger.error(f"Опрос: смена №{shift_number} не найдена в shifts_history.json (всего смен: {len(self.shifts)})")
                await event.message.answer(
                    "❌ <b>Не удалось сохранить ответ</b>\n\nЭта смена больше не найдена в базе бота "
                    "(вероятно, бот перезапускался и потерял данные). Запустите опрос заново через "
                    "<code>/test_survey</code> или дождитесь субботы.",
                    format=TextFormat.HTML
                )
                return

            question = SURVEY_CONFIG["questions"][question_key]
            answer_text = question["options"][option_index]
            shift["survey"][question_key] = answer_text
            self.save_shifts()

            responder = event.callback.user
            audit(str(responder.user_id), responder.username, "survey_answer",
                  f"смена №{shift_number}, {question_key}={answer_text}")

            order = ["quality", "incidents", "zgd"]
            next_index = order.index(question_key) + 1 if question_key in order else len(order)

            if next_index < len(order):
                next_key = order[next_index]
                next_question = SURVEY_CONFIG["questions"][next_key]
                kb = InlineKeyboardBuilder()
                for i, option in enumerate(next_question["options"]):
                    kb.row(CallbackButton(text=option, payload=f"survey|{shift_number}|{next_key}|{i}"))
                await event.message.edit(
                    text=f"✅ Ответ сохранён: <b>{answer_text}</b>\n\n{next_question['text']}",
                    attachments=[kb.as_markup()],
                    format=TextFormat.HTML
                )
            else:
                await event.message.edit(
                    text=f"✅ Ответ сохранён: <b>{answer_text}</b>\n\n"
                         f"📝 Последний шаг — напишите замечания текстом (или отправьте <code>-</code>, если их нет):",
                    attachments=[],  # явно убираем кнопки Да/Нет — иначе message.edit() без attachments
                                     # оставляет предыдущую клавиатуру видимой (она уже ни на что не влияет)
                    format=TextFormat.HTML
                )
                await context.update_data(survey_shift_number=shift_number)
                await context.set_state(SurveyWizard.awaiting_remarks)
        except Exception as e:
            logger.error(f"Ошибка обработки ответа на опрос (payload={payload}): {e}")
            await event.message.answer(f"❌ Ошибка сохранения ответа: {str(e)[:200]}", format=TextFormat.HTML)

    async def wizard_survey_remarks(self, event: MessageCreated, context: BaseContext):
        data = await context.get_data()
        shift_number = data.get("survey_shift_number")
        await context.clear()

        shift = self._get_shift(shift_number) if shift_number else None
        if not shift:
            await event.message.answer("❌ Не удалось найти смену для этого опроса.", format=TextFormat.HTML)
            return

        remarks = (event.message.body.text or "").strip()
        shift["survey"]["remarks"] = "" if remarks == "-" else remarks
        shift["completed"] = True
        self.save_shifts()

        sender = event.message.sender
        audit(str(sender.user_id), sender.username, "survey_completed", f"смена №{shift_number}")

        await event.message.answer(
            "✅ <b>ОПРОС ЗАВЕРШЁН</b>\n\nСпасибо! Протокол смены будет сформирован автоматически.",
            format=TextFormat.HTML
        )

        await self._finalize_shift_protocol(shift, event.message.recipient.chat_id)

    async def _finalize_shift_protocol(self, shift: Dict, chat_id: Optional[int]):
        """Генерирует .docx протокол по итогам смены и закрепляет его в чате (ТЗ п.2.5)."""
        try:
            date_compact = shift["date"].replace("г.", "").replace(".", "-")
            filename = PROTOCOL_FILENAME_MASK.format(date=date_compact, shift_number=shift["shift_number"])
            self.protocol_dir.mkdir(exist_ok=True)
            file_path = self.protocol_dir / filename

            doc = Document()
            doc.add_heading("Протокол разногласий", level=1)
            doc.add_paragraph(f"Дата дежурства: {shift['date']}")
            doc.add_paragraph(f"Смена №{shift['shift_number']}")
            doc.add_paragraph(f"Дежурили: {', '.join(shift['employees'])}")
            doc.add_heading("Итоги опроса", level=2)
            survey = shift.get("survey", {})
            doc.add_paragraph(f"Как прошло дежурство: {survey.get('quality', '—')}")
            doc.add_paragraph(f"Инциденты: {survey.get('incidents', '—')}")
            doc.add_paragraph(f"ЗГД: {survey.get('zgd', '—')}")
            doc.add_paragraph(f"Замечания: {survey.get('remarks') or '—'}")
            doc.save(str(file_path))

            shift["protocol_file"] = str(file_path)
            self.save_shifts()

            if chat_id is not None:
                sent = await self.bot.send_message(
                    chat_id=chat_id,
                    text=f"📄 Автопротокол смены №{shift['shift_number']} от {shift['date']}",
                    attachments=[InputMedia(path=str(file_path))],
                )
                mid = sent.message.body.mid if sent and sent.message and sent.message.body else None
                if mid:
                    await self.bot.pin_message(chat_id=chat_id, message_id=mid)

            logger.info(f"Автопротокол смены №{shift['shift_number']} сформирован: {file_path}")
            audit("system", None, "protocol_generated", f"смена №{shift['shift_number']} -> {filename}")
        except Exception as e:
            logger.error(f"Ошибка формирования автопротокола: {e}")

    async def _send_notification_to_all_users(self, message: str, notification_type: str):
        """Отправка уведомлений ВСЕМ пользователям с проверкой ID.

        NB: список фраз для распознавания заблокированных/недоступных
        пользователей был подобран под тексты ошибок Telegram Bot API.
        Формулировки ошибок MAX могут отличаться — стоит свериться с
        реальными логами после первого прогона и скорректировать список.
        """
        sent_count = 0
        error_count = 0
        deactivated_users = []

        self.load_user_data()
        logger.info(f"Отправка уведомления {notification_type} - всего пользователей: {len(self.user_data)}")

        for user_id, user_info in list(self.user_data.items()):
            try:
                await self.bot.send_message(user_id=int(user_id), text=message, format=TextFormat.HTML)
                sent_count += 1
                logger.debug(f"✓ Отправлено пользователю {user_id}")
                await asyncio.sleep(0.1)
            except ValueError:
                logger.error(f"✗ Некорректный ID пользователя: {user_id}")
                error_count += 1
                deactivated_users.append(user_id)
            except Exception as e:
                error_count += 1
                error_msg = str(e).lower()
                logger.error(f"✗ Ошибка отправки пользователю {user_id}: {error_msg[:100]}")

                if any(phrase in error_msg for phrase in [
                    'blocked', 'not found', 'kicked', 'deactivated', 'forbidden', 'access denied'
                ]):
                    logger.warning(f"Удаляю неактивного пользователя: {user_id}")
                    deactivated_users.append(user_id)

        for user_id in deactivated_users:
            self.user_data.pop(user_id, None)

        if deactivated_users:
            self.save_user_data()

        logger.info(f"=== ИТОГИ УВЕДОМЛЕНИЯ {notification_type.upper()} ===")
        logger.info(f"Всего в базе: {len(self.user_data) + len(deactivated_users)}")
        logger.info(f"Отправлено успешно: {sent_count}")
        logger.info(f"Ошибок: {error_count}")
        logger.info(f"Удалено неактивных: {len(deactivated_users)}")

        if sent_count == 0 and len(self.user_data) > 0:
            logger.error("⚠️ КРИТИЧЕСКАЯ ПРОБЛЕМА: НЕ УДАЛОСЬ ОТПРАВИТЬ НИ ОДНОГО УВЕДОМЛЕНИЯ!")

    # ================= ХРАНЕНИЕ ДАННЫХ ПОЛЬЗОВАТЕЛЕЙ =================

    def load_user_data(self):
        if os.path.exists(self.user_data_file):
            try:
                with open(self.user_data_file, 'r', encoding='utf-8') as f:
                    self.user_data = json.load(f)
            except Exception as e:
                logger.error(f"Ошибка загрузки user_data.json: {e}")
                self.user_data = {}
        else:
            self.user_data = {}

    def save_user_data(self):
        try:
            with open(self.user_data_file, 'w', encoding='utf-8') as f:
                json.dump(self.user_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Ошибка сохранения user_data.json: {e}")

    def load_shifts(self):
        """История дежурств для /stats и опроса по итогам смены (ТЗ п.2.1, 2.2).

        Каждая запись: {shift_number, date (дд.мм.гггг), employees, survey:
        {quality, incidents, zgd, remarks} | None, protocol_file, completed}.
        Статистика считается только по сменам, начиная с включения этой фичи —
        задним числом прошлые дежурства бот не восстанавливает."""
        if os.path.exists(self.shifts_file):
            try:
                with open(self.shifts_file, 'r', encoding='utf-8') as f:
                    self.shifts = json.load(f)
            except Exception as e:
                logger.error(f"Ошибка загрузки {self.shifts_file}: {e}")
                self.shifts = []
        else:
            self.shifts = []

    def save_shifts(self):
        try:
            with open(self.shifts_file, 'w', encoding='utf-8') as f:
                json.dump(self.shifts, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Ошибка сохранения {self.shifts_file}: {e}")

    def is_admin(self, user_id: str) -> bool:
        return self.user_data.get(user_id, {}).get("is_admin", False)

    def is_authorized_admin(self, user_id: str) -> bool:
        """Полная проверка админ-доступа (ТЗ п.3, «Безопасность»): либо
        user_id есть в ADMIN_IDS из конфига (доступ без /admin login),
        либо пользователь вошёл обычным логином/паролем.

        Статус входа хранится в user_data.json (persisted), а не только в
        памяти — иначе он слетал бы при каждом рестарте/redeploy бота на
        хостинге, вынуждая логиниться заново после буквально любого деплоя."""
        try:
            if int(user_id) in ADMIN_IDS:
                return True
        except ValueError:
            pass
        return self.is_admin(user_id) and self.user_data.get(user_id, {}).get("admin_logged_in", False)

    def is_whitelisted(self, user_id: str) -> bool:
        """Защита от посторонних (ТЗ п.2.4). Если WHITELIST_MODE выключен в
        конфиге — доступ открыт всем, как и было раньше. Сотрудники, автоматически
        привязанные по username из config.json, и админы считаются доверенными сразу."""
        if not WHITELIST_MODE:
            return True
        if self.is_admin(user_id) or self.is_authorized_admin(user_id):
            return True
        info = self.user_data.get(user_id, {})
        return bool(info.get("whitelisted") or info.get("selected_employee"))

    def get_employee_by_username(self, username: str) -> Optional[str]:
        if not username.startswith('@'):
            username = '@' + username
        return USERNAME_TO_EMPLOYEE.get(username.lower())

    # ================= КЛАВИАТУРЫ =================

    def get_main_keyboard(self, user_id: str):
        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="📋 Полный график", payload="full_schedule"),
            CallbackButton(text="👤 Моё дежурство", payload="my_duty"),
        )
        kb.row(
            CallbackButton(text="📄 Скачать протокол", payload="protocol"),
            CallbackButton(text="❓ Частые вопросы", payload="questions"),
        )
        kb.row(
            CallbackButton(text="📝 Инструкция", payload="instructions"),
            CallbackButton(text="🔄 Изменить профиль", payload="change_profile"),
        )
        kb.row(CallbackButton(text="📊 Статистика дежурств", payload="stats_menu"))
        if self.is_admin(user_id):
            kb.row(CallbackButton(text="⚙️ Админ-панель", payload="admin_panel"))
        return kb.as_markup()

    def get_admin_keyboard(self):
        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="📅 Управление графиком", payload="admin_schedule"),
            CallbackButton(text="👥 Управление сотрудниками", payload="admin_employees"),
        )
        kb.row(
            CallbackButton(text="📁 Управление файлами", payload="admin_files"),
            CallbackButton(text="📊 Статистика", payload="admin_stats"),
        )
        kb.row(
            CallbackButton(text="🔙 В главное меню", payload="back_to_main"),
            CallbackButton(text="🚪 Выйти из админки", payload="admin_logout"),
        )
        return kb.as_markup()

    def get_schedule_admin_keyboard(self):
        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="➕ Добавить дежурство", payload="admin_add_duty"),
            CallbackButton(text="➖ Удалить дежурство", payload="admin_remove_duty"),
        )
        kb.row(
            CallbackButton(text="📋 Просмотреть график", payload="full_schedule"),
            CallbackButton(text="🔄 Обновить график", payload="admin_refresh_schedule"),
        )
        kb.row(CallbackButton(text="🔙 Назад в админку", payload="admin_panel"))
        return kb.as_markup()

    def get_employees_admin_keyboard(self):
        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="➕ Добавить сотрудника", payload="admin_add_employee"),
            CallbackButton(text="➖ Удалить сотрудника", payload="admin_remove_employee"),
        )
        kb.row(
            CallbackButton(text="📞 Изменить телефон", payload="admin_edit_phone"),
            CallbackButton(text="👥 Список сотрудников", payload="admin_list_employees"),
        )
        kb.row(CallbackButton(text="🔙 Назад в админку", payload="admin_panel"))
        return kb.as_markup()

    def get_files_admin_keyboard(self):
        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="📤 Загрузить протокол", payload="admin_upload_protocol"),
            CallbackButton(text="📎 Прикрепить протокол", payload="admin_pin_protocol"),
        )
        kb.row(
            CallbackButton(text="🗑 Удалить протокол", payload="admin_delete_protocol"),
            CallbackButton(text="📄 Проверить файл", payload="admin_check_protocol"),
        )
        kb.row(CallbackButton(text="🔙 Назад в админку", payload="admin_panel"))
        return kb.as_markup()

    def get_back_keyboard(self):
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="🔙 Назад в меню", payload="back_to_main"))
        return kb.as_markup()

    def get_employee_selection_keyboard(self, prefix: str = "emp_"):
        kb = InlineKeyboardBuilder()
        employees_list = list(EMPLOYEE_PHONES.keys())

        for i in range(0, len(employees_list), 2):
            row_buttons = [CallbackButton(text=employees_list[i], payload=f"{prefix}{employees_list[i]}")]
            if i + 1 < len(employees_list):
                row_buttons.append(
                    CallbackButton(text=employees_list[i + 1], payload=f"{prefix}{employees_list[i + 1]}"))
            kb.row(*row_buttons)

        if prefix.startswith("add_e"):
            kb.row(CallbackButton(text="❌ Отмена", payload="admin_schedule"))

        return kb.as_markup()

    # ================= КОМАНДЫ =================

    @staticmethod
    def _parse_args(event: MessageCreated) -> List[str]:
        text = (event.message.body.text or "") if event.message.body else ""
        parts = text.split()
        return parts[1:] if len(parts) > 1 else []

    def _register_user_and_build_welcome(self, user_id: str, username: Optional[str],
                                          first_name: Optional[str], last_name: Optional[str]):
        """Регистрирует пользователя (если это первый заход) и собирает текст+клавиатуру
        приветственного меню. Используется и командой /start, и авто-показом меню при
        открытии чата с ботом (BotStarted) — единая точка правды для обоих путей."""
        if user_id not in self.user_data:
            self.user_data[user_id] = {
                "username": username,
                "first_name": first_name,
                "last_name": last_name,
                "display_name": f"{first_name or ''} {last_name or ''}".strip(),
                "notifications": True,
                "selected_employee": None,
                "registered_at": datetime.now().isoformat(),
                "last_active": datetime.now().isoformat(),
                "is_admin": False
            }

            if username:
                employee_name = self.get_employee_by_username(username)
                if employee_name:
                    self.user_data[user_id]["selected_employee"] = employee_name

            self.save_user_data()

        self.user_data[user_id]["last_active"] = datetime.now().isoformat()
        self.save_user_data()

        user_info = self.user_data[user_id]
        employee_name = user_info.get("selected_employee")

        if not self.is_whitelisted(user_id):
            status = user_info.get("access_request_status")
            if status == "pending":
                text = (
                    f"<b>ДОБРО ПОЖАЛОВАТЬ, {first_name}!</b>\n\n"
                    "⏳ Ваш запрос на доступ уже отправлен администратору и ожидает подтверждения."
                )
                kb = None
            else:
                text = (
                    f"<b>ДОБРО ПОЖАЛОВАТЬ, {first_name}!</b>\n\n"
                    "🔒 <b>Доступ ограничен.</b>\n\n"
                    "Этот бот доступен только сотрудникам. Если вы сотрудник — "
                    "запросите доступ у администратора."
                )
                builder = InlineKeyboardBuilder()
                builder.row(CallbackButton(text="🔓 Запросить доступ", payload="request_access"))
                kb = builder.as_markup()
            return text, kb

        if employee_name:
            welcome_text = (
                f"<b>ДОБРО ПОЖАЛОВАТЬ, {first_name}!</b>\n\n"
                f"👤 <b>Ваш профиль:</b>\n"
                f"• Сотрудник: {employee_name}\n"
                f"• Телефон: {EMPLOYEE_PHONES.get(employee_name, 'не указан')}\n"
                f"• Уведомления: {'✅ Включены' if user_info.get('notifications', True) else '❌ Отключены'}\n\n"
                "<i>Выберите действие:</i>"
            )
            keyboard = self.get_main_keyboard(user_id)
        else:
            welcome_text = (
                f"<b>ДОБРО ПОЖАЛОВАТЬ, {first_name}!</b>\n\n"
                "Я бот для управления графиком дежурств.\n\n"
            )
            if username:
                welcome_text += f"Ваш username: @{username}\n"
            welcome_text += "<i>Пожалуйста, выберите ваше ФИО из списка:</i>"
            keyboard = self.get_employee_selection_keyboard(prefix="emp_")

        return welcome_text, keyboard

    async def start(self, event: MessageCreated):
        sender = event.message.sender
        user_id = str(sender.user_id)
        welcome_text, keyboard = self._register_user_and_build_welcome(
            user_id, sender.username, sender.first_name, sender.last_name
        )
        await event.message.answer(
            welcome_text,
            attachments=[keyboard] if keyboard else [],
            format=TextFormat.HTML
        )

    async def on_bot_started(self, event: BotStarted):
        """Срабатывает, когда пользователь открывает чат с ботом впервые —
        аналог 'кнопки меню': показываем меню сразу, без ожидания команды /start."""
        user = event.user
        user_id = str(user.user_id)
        welcome_text, keyboard = self._register_user_and_build_welcome(
            user_id, user.username, user.first_name, user.last_name
        )
        await self.bot.send_message(
            user_id=user.user_id,
            text=welcome_text,
            attachments=[keyboard] if keyboard else [],
            format=TextFormat.HTML
        )

    async def admin_login(self, event: MessageCreated):
        sender = event.message.sender
        user_id = str(sender.user_id)
        args = self._parse_args(event)

        if len(args) != 2:
            await event.message.answer(
                "❌ <b>Неверный формат команды</b>\n\nИспользуйте: /admin логин пароль\n",
                format=TextFormat.HTML
            )
            return

        login, password = args[0], args[1]

        if login == ADMIN_CREDENTIALS["login"] and password == ADMIN_CREDENTIALS["password"]:
            if user_id not in self.user_data:
                # Логин без предварительного /start — заводим карточку пользователя,
                # иначе is_admin/admin_logged_in было бы некуда сохранить.
                self.user_data[user_id] = {
                    "username": sender.username, "first_name": sender.first_name,
                    "last_name": sender.last_name,
                    "display_name": f"{sender.first_name or ''} {sender.last_name or ''}".strip(),
                    "notifications": True, "selected_employee": None,
                    "registered_at": datetime.now().isoformat(),
                    "last_active": datetime.now().isoformat(), "is_admin": False,
                }

            self.user_data[user_id]["is_admin"] = True
            self.user_data[user_id]["admin_logged_in"] = True
            self.user_data[user_id]["admin_login_time"] = datetime.now().isoformat()
            self.save_user_data()
            audit(user_id, sender.username, "admin_login_success")

            await event.message.answer(
                "✅ <b>УСПЕШНЫЙ ВХОД В АДМИН-ПАНЕЛЬ</b>\n\n"
                "Доступные функции:\n"
                "• Управление графиком дежурств\n"
                "• Управление сотрудниками\n"
                "• Управление файлами\n"
                "• Просмотр статистики\n\n"
                "<i>Выберите действие:</i>",
                attachments=[self.get_admin_keyboard()],
                format=TextFormat.HTML
            )
        else:
            audit(user_id, sender.username, "admin_login_failed")
            await event.message.answer(
                "❌ <b>НЕВЕРНЫЙ ЛОГИН ИЛИ ПАРОЛЬ</b>\n\nПопробуйте снова:\n",
                format=TextFormat.HTML
            )

    # ============= ДИАГНОСТИЧЕСКИЕ КОМАНДЫ ДЛЯ СУПЕР-АДМИНА =============

    async def check_users_status(self, event: MessageCreated):
        sender = event.message.sender
        if not self.is_authorized_admin(str(sender.user_id)):
            await event.message.answer(
                "❌ <b>ДОСТУП ЗАПРЕЩЕН</b>\n\nЭта команда доступна только админам, вошедшим через /admin",
                format=TextFormat.HTML
            )
            return

        self.load_user_data()
        text = "📊 <b>СТАТУС ПОЛЬЗОВАТЕЛЕЙ</b>\n\n"
        total = len(self.user_data)
        with_employee = 0
        notifications_on = 0
        notifications_off = 0

        for uid, info in self.user_data.items():
            name = info.get('display_name', 'Неизвестно')
            username = info.get('username', 'Нет username')
            employee = info.get('selected_employee', None)
            notifications = info.get('notifications', True)

            if employee and employee != 'None' and employee != '❌ НЕ ВЫБРАН':
                with_employee += 1
            if notifications:
                notifications_on += 1
            else:
                notifications_off += 1

            notif_status = "✅ ВКЛ" if notifications else "❌ ВЫКЛ"
            employee_display = employee if employee else "❌ НЕ ВЫБРАН"

            text += f"<b>{name}</b>\n📱 @{username}\n🆔 {uid}\n👤 {employee_display}\n🔔 {notif_status}\n📅 Последний вход: {info.get('last_active', 'Неизвестно')[:16]}\n\n"

        text += f"<b>ИТОГО:</b> {total} пользователей\n👤 С выбором сотрудника: {with_employee}\n🔔 Уведомления включены: {notifications_on}\n🔕 Уведомления выключены: {notifications_off}"
        await event.message.answer(text, format=TextFormat.HTML)

    async def enable_notifications_all(self, event: MessageCreated):
        sender = event.message.sender
        if not self.is_authorized_admin(str(sender.user_id)):
            return

        self.load_user_data()
        enabled_count = 0
        for uid, info in self.user_data.items():
            if not info.get('notifications', True):
                self.user_data[uid]['notifications'] = True
                enabled_count += 1

        self.save_user_data()
        await event.message.answer(
            f"✅ Уведомления включены для {enabled_count} пользователей\n📊 Всего пользователей: {len(self.user_data)}",
            format=TextFormat.HTML)

    async def test_send_to_user(self, event: MessageCreated):
        sender = event.message.sender
        if not self.is_authorized_admin(str(sender.user_id)):
            return

        args = self._parse_args(event)
        if len(args) < 1:
            await event.message.answer(
                "❌ Укажите user_id или username\nПример: /test_send 123456789\nИли: /test_send @username")
            return

        target = args[0]
        target_id = None
        target_name = target

        if target.startswith('@'):
            username = target[1:].lower()
            for uid, info in self.user_data.items():
                if info.get('username', '').lower() == username:
                    target_id = uid
                    target_name = info.get('display_name', target)
                    break
            if not target_id:
                await event.message.answer(f"❌ Пользователь {target} не найден в базе")
                return
        else:
            target_id = target

        test_msg = (
            f"🔔 <b>ТЕСТОВОЕ УВЕДОМЛЕНИЕ</b>\n\n👤 Получатель: {target_name}\n🆔 ID: {target_id}\n📅 Время: {datetime.now(MOSCOW_TZ).strftime('%d.%m.%Y %H:%M:%S')}\n\n"
            f"✅ Если вы видите это сообщение, значит:\n   • Бот может отправлять вам сообщения\n   • Вы не блокировали бота\n   • Уведомления будут приходить по расписанию\n\n"
            f"📅 Расписание уведомлений:\n• Среда 18:00 - о дежурстве в субботу\n• Пятница 18:00 - о завтрашнем дежурстве\n• Суббота 10:00 - в день дежурства"
        )

        try:
            await self.bot.send_message(user_id=int(target_id), text=test_msg, format=TextFormat.HTML)
            await event.message.answer(f"✅ Тестовое сообщение отправлено {target_name}")
        except Exception as e:
            await event.message.answer(f"❌ Ошибка отправки: {str(e)[:200]}")

    async def check_time(self, event: MessageCreated):
        sender = event.message.sender
        if not self.is_authorized_admin(str(sender.user_id)):
            return

        now = datetime.now(MOSCOW_TZ)
        weekdays = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]
        weekday_ru = weekdays[now.weekday()]

        next_notification = ""
        if now.weekday() == 1 and now.hour < 18:
            next_notification = "Среда 18:00 (через 1 день)"
        elif now.weekday() == 2 and now.hour < 18:
            next_notification = "Среда 18:00 (сегодня)"
        elif now.weekday() == 3:
            next_notification = "Пятница 18:00 (через 1 день)"
        elif now.weekday() == 4 and now.hour < 18:
            next_notification = "Пятница 18:00 (сегодня)"
        elif now.weekday() == 5 and now.hour < 10:
            next_notification = "Суббота 10:00 (сегодня)"
        elif now.weekday() == 6:
            next_notification = "Среда 18:00 (через 3 дня)"
        else:
            next_notification = "Среда 18:00"

        await event.message.answer(
            f"🕐 <b>ИНФОРМАЦИЯ О ВРЕМЕНИ</b>\n\n📅 Дата: {now.strftime('%d.%m.%Y')}\n⏰ Время: {now.strftime('%H:%M:%S')}\n📆 День недели: {weekday_ru}\n🌍 Часовой пояс: Москва (UTC+3)\n\n"
            f"🔄 <b>Следующее уведомление:</b> {next_notification}\n\n📋 <b>Расписание:</b>\n• Среда 18:00 - всем\n• Пятница 18:00 - всем\n• Суббота 10:00 - всем\n\n"
            f"<i>Версия кода: {BOT_CODE_VERSION}</i>",
            format=TextFormat.HTML
        )

    async def fix_all_users(self, event: MessageCreated):
        sender = event.message.sender
        if not self.is_authorized_admin(str(sender.user_id)):
            return

        self.load_user_data()
        fixed_count = 0
        for uid, info in self.user_data.items():
            changes = []
            if not info.get('notifications', True):
                info['notifications'] = True
                changes.append("включены уведомления")
            if 'display_name' not in info:
                info['display_name'] = info.get('first_name', 'Пользователь')
                changes.append("добавлено имя")
            if changes:
                fixed_count += 1
                logger.info(f"Исправлен пользователь {uid}: {', '.join(changes)}")

        self.save_user_data()

        test_msg = (
            f"🔔 <b>ТЕСТОВОЕ УВЕДОМЛЕНИЕ ОТ АДМИНИСТРАТОРА</b>\n\n✅ Ваши уведомления были включены!\n\n"
            f"📅 Вы будете получать напоминания:\n• В среду в 18:00 - о дежурстве в субботу\n• В пятницу в 18:00 - о завтрашнем дежурстве\n• В субботу в 10:00 - в день дежурства\n\n📋 Используйте /start для просмотра меню"
        )

        sent_count = 0
        error_count = 0
        for uid in self.user_data.keys():
            try:
                await self.bot.send_message(user_id=int(uid), text=test_msg, format=TextFormat.HTML)
                sent_count += 1
                await asyncio.sleep(0.1)
            except Exception as e:
                error_count += 1
                logger.error(f"Ошибка отправки теста пользователю {uid}: {e}")

        await event.message.answer(
            f"✅ <b>ИСПРАВЛЕНИЕ ЗАВЕРШЕНО</b>\n\n📊 Исправлено пользователей: {fixed_count}\n📤 Отправлено тестовых уведомлений: {sent_count}\n❌ Ошибок отправки: {error_count}\n\n🔔 Теперь все пользователи будут получать уведомления!",
            format=TextFormat.HTML
        )

    # ============= КОНЕЦ ДИАГНОСТИЧЕСКИХ КОМАНД =============

    async def send_test_wednesday(self, event: MessageCreated):
        user_id = str(event.message.sender.user_id)
        if not self.is_admin(user_id):
            return
        await event.message.answer("🔄 Отправляю тестовое среднее уведомление всем пользователям...")
        await self.send_wednesday_notification()
        await event.message.answer("✅ Тестовое среднее уведомление отправлено!")

    async def send_test_friday(self, event: MessageCreated):
        user_id = str(event.message.sender.user_id)
        if not self.is_admin(user_id):
            return
        await event.message.answer("🔄 Отправляю тестовое пятничное уведомление всем пользователям...")
        await self.send_friday_notification_all()
        await event.message.answer("✅ Тестовое пятничное уведомление отправлено!")

    async def send_test_saturday(self, event: MessageCreated):
        user_id = str(event.message.sender.user_id)
        if not self.is_admin(user_id):
            return
        await event.message.answer("🔄 Отправляю тестовое субботнее уведомление всем пользователям...")
        await self.send_saturday_notification_all()
        await event.message.answer("✅ Тестовое субботнее уведомление отправлено!")

    async def send_test_survey(self, event: MessageCreated):
        """/test_survey — ручной запуск опроса по дежурству для тестирования
        в любой день недели (force=True снимает проверку «только суббота»)."""
        user_id = str(event.message.sender.user_id)
        if not self.is_admin(user_id):
            return
        await event.message.answer("🔄 Отправляю тестовый опрос по итогам дежурства...")
        result = await self.send_shift_survey(force=True)

        if not result["ok"]:
            await event.message.answer(f"❌ Опрос не отправлен: {result['reason']}", format=TextFormat.HTML)
            return

        text = f"Дежурные по графику: {', '.join(result['employees'])}\n📤 Реально получили опрос: {result['sent']}"
        if result.get("unlinked"):
            text += (
                f"\n\n⚠️ <b>Не привязаны ни к одному пользователю бота:</b> "
                f"{', '.join(result['unlinked'])}\n"
                f"Опрос им отправить некуда — сотрудник должен сначала написать /start "
                f"и выбрать своё ФИО в меню."
            )
        if result["sent"] == 0:
            text = "⚠️ Опрос создан, но <b>не ушёл ни одному получателю</b>.\n\n" + text

        await event.message.answer(text, format=TextFormat.HTML)

    async def test_notification_for_user(self, event: MessageCreated):
        user_id = str(event.message.sender.user_id)
        if not self.is_admin(user_id):
            return

        args = self._parse_args(event)
        if len(args) != 1:
            await event.message.answer("❌ Неверный формат. Используйте: /test_user <user_id>")
            return

        target_user_id = args[0]
        test_message = (
            f"🔔 <b>ТЕСТОВОЕ УВЕДОМЛЕНИЕ</b>\n\n📅 <b>Это тестовое сообщение от администратора</b>\n\n... "
        )
        try:
            await self.bot.send_message(user_id=int(target_user_id), text=test_message, format=TextFormat.HTML)
            await event.message.answer(f"✅ Тестовое сообщение отправлено пользователю {target_user_id}")
        except Exception as e:
            await event.message.answer(f"❌ Ошибка отправки: {str(e)}")

    # ================= ОБРАБОТКА КНОПОК (CALLBACK) =================

    async def on_callback(self, event: MessageCallback, context: BaseContext):
        if event.callback is None or event.message is None:
            if event.callback:
                await event.answer()
            return

        payload = event.callback.payload or ""
        user_id = str(event.callback.user.user_id)
        message = event.message

        await event.answer()

        if payload == "admin_panel":
            if self.is_authorized_admin(user_id):
                await self.show_admin_panel(message, user_id, context)
            else:
                await message.edit(
                    text="❌ <b>ДОСТУП ЗАПРЕЩЕН</b>\n\nДоступ только админам\n<code>Зайдите с нужного аккаунта!!</code>",
                    format=TextFormat.HTML)
            return

        handlers = {
            "full_schedule": self.show_full_schedule,
            "my_duty": self.show_my_duty,
            "instructions": self.show_instructions,
            "protocol": self.download_protocol,
            "questions": self.show_questions,
            "back_to_main": self.back_to_main,
            "change_profile": self.change_profile,
            "admin_logout": self.admin_logout,
            "admin_refresh_schedule": self.admin_refresh_schedule,
            "admin_schedule": self.show_admin_schedule,
            "admin_employees": self.show_admin_employees,
            "admin_files": self.show_admin_files,
            "admin_stats": self.show_admin_stats,
            "admin_remove_duty": self.admin_remove_duty,
            "admin_add_employee": self.admin_add_employee,
            "admin_remove_employee": self.admin_remove_employee,
            "admin_edit_phone": self.admin_edit_phone,
            "admin_list_employees": self.admin_list_employees,
            "admin_upload_protocol": self.admin_upload_protocol,
            "admin_delete_protocol": self.admin_delete_protocol,
            "admin_check_protocol": self.admin_check_protocol,
            "admin_pin_protocol": self.admin_pin_protocol,
            "stats_menu": self.show_stats_menu,
            "stats_week": self.show_stats_week,
            "stats_month": self.show_stats_month,
            "stats_all": self.show_stats_all,
            "request_access": self.request_access,
        }

        if payload.startswith("emp_"):
            employee_name = payload[4:]
            await self.register_employee(message, user_id, employee_name)

        elif payload == "admin_add_duty":
            kb = InlineKeyboardBuilder()
            kb.row(
                CallbackButton(text="👤 Один дежурный", payload="add_type_single"),
                CallbackButton(text="👥 Пара (2 чел.)", payload="add_type_pair"),
            )
            kb.row(CallbackButton(text="❌ Отмена", payload="admin_schedule"))
            await message.edit(
                text="➕ <b>НОВОЕ ДЕЖУРСТВО</b>\n\n<b>Шаг 1:</b> Выберите формат дежурства:",
                attachments=[kb.as_markup()],
                format=TextFormat.HTML
            )

        elif payload.startswith("add_type_"):
            is_pair = (payload == "add_type_pair")
            await context.update_data(new_duty={"is_pair": is_pair, "employees": [], "phones": []})
            await context.set_state(AdminWizard.wait_date)
            await message.edit(
                text="📅 <b>Шаг 2:</b> Введите дату в чат\n\nФормат: <code>дд.мм.гггг</code>\n<i>Например: 07.02.2026</i>",
                format=TextFormat.HTML
            )

        elif payload.startswith("add_e1_"):
            name = payload.replace("add_e1_", "")
            data = await context.get_data()
            duty = data.get("new_duty", {"is_pair": False, "employees": [], "phones": []})
            duty["employees"].append(name)
            await context.update_data(new_duty=duty)

            if duty["is_pair"]:
                await message.edit(
                    text=f"✅ Выбран первый: <b>{name}</b>\n\n👥 <b>Шаг 4:</b> Выберите второго дежурного:",
                    attachments=[self.get_employee_selection_keyboard("add_e2_")],
                    format=TextFormat.HTML
                )
            else:
                await context.set_state(AdminWizard.wait_phones)
                await message.edit(
                    text=f"✅ Выбран: <b>{name}</b>\n\n📞 <b>Шаг 4:</b> Введите номер телефона в чат.\n\n<i>Лайфхак: Напишите в чат слово <b>ок</b>, и бот сам подставит сохраненный номер сотрудника!</i>",
                    format=TextFormat.HTML
                )

        elif payload.startswith("add_e2_"):
            name = payload.replace("add_e2_", "")
            data = await context.get_data()
            duty = data["new_duty"]
            duty["employees"].append(name)
            await context.update_data(new_duty=duty)
            await context.set_state(AdminWizard.wait_phones)
            await message.edit(
                text=(f"✅ Выбрана пара: <b>{duty['employees'][0]} + {name}</b>\n\n"
                      "📞 <b>Шаг 5:</b> Введите телефоны через запятую.\n\n<i>Лайфхак: Напишите в чат слово <b>ок</b>, и бот сам подставит номера обоих сотрудников!</i>"),
                format=TextFormat.HTML
            )

        elif payload.startswith("survey|"):
            await self.handle_survey_callback(event, context)

        elif payload.startswith("access_approve|") or payload.startswith("access_deny|"):
            approved = payload.startswith("access_approve|")
            target_user_id = payload.split("|", 1)[1]
            await self.handle_access_decision(message, user_id, target_user_id, approved)

        elif payload in handlers:
            await handlers[payload](message, user_id, context)

    # ================= ЭКРАНЫ (вызываются из on_callback) =================

    async def show_full_schedule(self, message, user_id, context=None):
        text = self.schedule_generator.get_schedule_text()
        if len(text) > 4000:
            await message.edit(text=text[:4000], format=TextFormat.HTML)
            await message.reply(text=text[4000:], attachments=[self.get_back_keyboard()], format=TextFormat.HTML)
        else:
            await message.edit(text=text, attachments=[self.get_back_keyboard()], format=TextFormat.HTML)

    async def show_my_duty(self, message, user_id, context=None):
        if user_id not in self.user_data:
            await message.edit(text="❌ Сначала зарегистрируйтесь /start", format=TextFormat.HTML)
            return

        employee_name = self.user_data[user_id].get("selected_employee")
        if not employee_name:
            await message.edit(text="❌ Выберите сотрудника в меню", format=TextFormat.HTML)
            return

        duties = self.schedule_generator.get_employee_schedule(employee_name)
        today = datetime.now(MOSCOW_TZ).replace(tzinfo=None)

        if not duties:
            text = f"📅 <b>БЛИЖАЙШИЕ ДЕЖУРСТВА: {employee_name}</b>\n\nНет запланированных дежурств"
        else:
            text = f"📅 <b>БЛИЖАЙШИЕ ДЕЖУРСТВА: {employee_name}</b>\n\n"
            for duty in duties[:3]:
                days_left = (duty["date_obj"] - today).days
                if duty["is_pair"]:
                    partners = [e for e in duty["employees"] if e != employee_name]
                    duty_text = f"{duty['date']} (с {', '.join(partners)})"
                    phones = ', '.join(duty['phones'])
                else:
                    duty_text = duty['date']
                    phones = duty['phones'][0]

                text += f"{duty_text}\n📅 Осталось: {max(0, days_left)} дней\n📞 {phones}\n\n"

        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="🔙 Назад", payload="back_to_main"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)

    # ================= СТАТИСТИКА ДЕЖУРСТВ (ТЗ п.2.1) =================
    # Доступна всем участникам без ограничений по ролям. Считает только смены,
    # завершённые ПОСЛЕ включения этой функции — прошлых дежурств бот не знает.

    def _build_stats_text(self, period: str) -> str:
        now = datetime.now(MOSCOW_TZ).replace(tzinfo=None)
        if period == "week":
            since = now - timedelta(days=7)
            title = "за последнюю неделю"
        elif period == "month":
            since = now - timedelta(days=30)
            title = "за последний месяц"
        else:
            since = None
            title = "за всё время"

        counts: Dict[str, int] = {}
        total = 0
        for shift in self.shifts:
            try:
                shift_date = datetime.strptime(shift["date"].replace("г.", "").strip(), "%d.%m.%Y")
            except Exception:
                continue
            if since and shift_date < since:
                continue
            for emp in shift.get("employees", []):
                counts[emp] = counts.get(emp, 0) + 1
                total += 1

        text = f"📊 <b>СТАТИСТИКА ДЕЖУРСТВ</b>\n<i>{title}</i>\n\n"
        if total == 0:
            text += "Нет завершённых дежурств за этот период.\n\n<i>Учёт ведётся с момента подключения этой функции.</i>"
        else:
            for emp, cnt in sorted(counts.items(), key=lambda x: -x[1]):
                pct = (cnt / total) * 100
                text += f"• <b>{emp}</b> — {cnt} ({pct:.0f}%)\n"
            text += f"\n<b>Всего дежурств за период:</b> {total}"

        return text

    def _get_stats_keyboard(self):
        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="Неделя", payload="stats_week"),
            CallbackButton(text="Месяц", payload="stats_month"),
            CallbackButton(text="Всё время", payload="stats_all"),
        )
        kb.row(CallbackButton(text="🔙 Назад", payload="back_to_main"))
        return kb.as_markup()

    async def show_stats_menu(self, message, user_id, context=None):
        await message.edit(
            text="📊 <b>СТАТИСТИКА ДЕЖУРСТВ</b>\n\nВыберите период:",
            attachments=[self._get_stats_keyboard()],
            format=TextFormat.HTML
        )

    async def show_stats_week(self, message, user_id, context=None):
        await message.edit(text=self._build_stats_text("week"), attachments=[self._get_stats_keyboard()],
                            format=TextFormat.HTML)

    async def show_stats_month(self, message, user_id, context=None):
        await message.edit(text=self._build_stats_text("month"), attachments=[self._get_stats_keyboard()],
                            format=TextFormat.HTML)

    async def show_stats_all(self, message, user_id, context=None):
        await message.edit(text=self._build_stats_text("all"), attachments=[self._get_stats_keyboard()],
                            format=TextFormat.HTML)

    async def cmd_stats(self, event: MessageCreated):
        """Команда /stats — дублирует кнопку «Статистика дежурств» (ТЗ п.5, UI/UX)."""
        await event.message.answer(
            "📊 <b>СТАТИСТИКА ДЕЖУРСТВ</b>\n\nВыберите период:",
            attachments=[self._get_stats_keyboard()],
            format=TextFormat.HTML
        )

    async def cmd_contact(self, event: MessageCreated):
        """Команда /contact — показывает привязанного сотрудника и его телефон."""
        user_id = str(event.message.sender.user_id)
        user_info = self.user_data.get(user_id, {})
        employee_name = user_info.get("selected_employee")
        if employee_name:
            text = (
                f"👤 <b>Ваш профиль</b>\n\n"
                f"Сотрудник: {employee_name}\n"
                f"Телефон: {EMPLOYEE_PHONES.get(employee_name, 'не указан')}"
            )
        else:
            text = "❌ У вас не выбран сотрудник. Откройте меню → «Изменить профиль»."
        await event.message.answer(text, format=TextFormat.HTML)

    async def cmd_set_phone(self, event: MessageCreated):
        """Команда /set_phone — алиас кнопки «Изменить телефон» (только для админов из ADMIN_IDS/логина).

        ВАЖНО: не переиспользует admin_edit_phone() напрямую — тот вызывает
        message.edit(), а бот не может редактировать чужое (не своё) сообщение,
        которым здесь как раз является входящая команда /set_phone."""
        user_id = str(event.message.sender.user_id)
        if not self.is_authorized_admin(user_id):
            await event.message.answer("❌ Команда доступна только администраторам.", format=TextFormat.HTML)
            return

        employees_list = "\n".join([f"• {emp}" for emp in EMPLOYEE_PHONES.keys()])
        text = (
            "📞 <b>ИЗМЕНЕНИЕ ТЕЛЕФОНА СОТРУДНИКА</b>\n\n"
            f"<b>Список сотрудников:</b>\n{employees_list}\n\n"
            "<i>Для изменения телефона отправьте сообщение в формате:</i>\n\n"
            "<code>ФИО;новый телефон</code>\n\n"
            "<b>Пример:</b>\n<code>Иванов И.И.;8-900-000-00-00</code>"
        )
        await event.message.answer(text, format=TextFormat.HTML)
        context = self.dp.fsm.get_context(chat_id=event.message.recipient.chat_id, user_id=int(user_id))
        await context.set_state(AdminWizard.awaiting_phone_edit)

    async def cmd_set_schedule(self, event: MessageCreated):
        """Команда /set_schedule — алиас кнопки «Управление графиком» (только для админов).

        По той же причине, что и cmd_set_phone, не вызывает show_admin_schedule()
        напрямую (там message.edit()) — вместо этого просто отправляет меню новым сообщением."""
        user_id = str(event.message.sender.user_id)
        if not self.is_authorized_admin(user_id):
            await event.message.answer("❌ Команда доступна только администраторам.", format=TextFormat.HTML)
            return

        text = (
            "📅 <b>УПРАВЛЕНИЕ ГРАФИКОМ ДЕЖУРСТВ</b>\n\n"
            "➕ Добавить дежурство / ➖ Удалить дежурство / 📋 Просмотреть график:"
        )
        await event.message.answer(text, attachments=[self.get_schedule_admin_keyboard()], format=TextFormat.HTML)

    async def cmd_myid(self, event: MessageCreated):
        """/myid — без проверки прав: показывает числовой ID отправителя,
        чтобы его можно было вписать в config.json -> admin_ids. Доступ по
        username больше нигде в боте не используется — только по ID/логину."""
        sender = event.message.sender
        await event.message.answer(
            f"🆔 Ваш ID: <code>{sender.user_id}</code>\n\n"
            f"Чтобы получить права админа без /admin login — добавьте это число "
            f"в config.json (или CONFIG_JSON) -> <code>admin_ids</code>.",
            format=TextFormat.HTML
        )

    # ================= ЗАЩИТА ОТ ПОСТОРОННИХ (ТЗ п.2.4) =================

    async def request_access(self, message, user_id, context=None):
        user_info = self.user_data.get(user_id, {})
        if user_info.get("access_request_status") == "pending":
            await message.edit(text="⏳ Запрос уже отправлен, ожидайте подтверждения.", format=TextFormat.HTML)
            return

        user_info["access_request_status"] = "pending"
        self.save_user_data()

        display_name = user_info.get("display_name") or "Без имени"
        username = user_info.get("username")

        kb = InlineKeyboardBuilder()
        kb.row(
            CallbackButton(text="✅ Разрешить", payload=f"access_approve|{user_id}"),
            CallbackButton(text="❌ Отклонить", payload=f"access_deny|{user_id}"),
        )
        notify_text = (
            f"🔔 <b>ЗАПРОС ДОСТУПА К БОТУ</b>\n\n"
            f"👤 {display_name}\n"
            f"📱 @{username or 'нет username'}\n"
            f"🆔 {user_id}"
        )

        admin_ids = [uid for uid, info in self.user_data.items() if info.get("is_admin")]
        for admin_id in admin_ids:
            try:
                await self.bot.send_message(user_id=int(admin_id), text=notify_text,
                                             attachments=[kb.as_markup()], format=TextFormat.HTML)
            except Exception as e:
                logger.error(f"Не удалось уведомить админа {admin_id} о запросе доступа: {e}")

        audit(user_id, username, "access_requested")
        await message.edit(text="✅ Запрос отправлен администратору. Ожидайте подтверждения.",
                            format=TextFormat.HTML)

    async def handle_access_decision(self, message, admin_user_id: str, target_user_id: str, approved: bool):
        if not self.is_authorized_admin(admin_user_id):
            await message.edit(text="❌ Только администратор может подтверждать доступ.", format=TextFormat.HTML)
            return

        target_info = self.user_data.get(target_user_id)
        if not target_info:
            await message.edit(text="❌ Пользователь не найден.", format=TextFormat.HTML)
            return

        target_name = target_info.get("display_name", target_user_id)
        if approved:
            target_info["whitelisted"] = True
            target_info["access_request_status"] = "approved"
            self.save_user_data()
            await message.edit(text=f"✅ Доступ разрешён для {target_name}.", attachments=[], format=TextFormat.HTML)
            notify_text = "✅ Ваш доступ к боту подтверждён администратором! Напишите /start."
        else:
            target_info["access_request_status"] = "denied"
            self.save_user_data()
            await message.edit(text=f"❌ Доступ отклонён для {target_name}.", attachments=[], format=TextFormat.HTML)
            notify_text = "❌ Ваш запрос на доступ к боту отклонён администратором."

        try:
            await self.bot.send_message(user_id=int(target_user_id), text=notify_text, format=TextFormat.HTML)
        except Exception as e:
            logger.error(f"Не удалось уведомить пользователя {target_user_id} о решении по доступу: {e}")

        audit(admin_user_id, None, "access_decision", f"target={target_user_id}, approved={approved}")

    async def back_to_main(self, message, user_id, context=None):
        user_info = self.user_data.get(user_id, {})
        employee_name = user_info.get("selected_employee")

        if employee_name:
            text = (
                "<b>🏠 ГЛАВНОЕ МЕНЮ</b>\n\n"
                f"👤 <b>Сотрудник:</b> {employee_name}\n"
                f"📞 <b>Телефон:</b> {EMPLOYEE_PHONES.get(employee_name, 'не указан')}\n\n"
                "<i>Выберите действие:</i>"
            )
        else:
            text = "<b>🏠 ГЛАВНОЕ МЕНЮ</b>\n\n<i>Для доступа к функциям\nнеобходима регистрация.</i>\n\nВыберите действие:"

        await message.edit(text=text, attachments=[self.get_main_keyboard(user_id)], format=TextFormat.HTML)

    async def download_protocol(self, message, user_id, context=None):
        try:
            if not os.path.exists(self.protocol_file_path):
                await message.edit(text="❌ Файл не найден", attachments=[self.get_back_keyboard()],
                                    format=TextFormat.HTML)
                return

            await message.reply(
                text="📄 Протокол разногласий",
                attachments=[InputMedia(path=self.protocol_file_path)],
            )
            await message.edit(text="✅ Файл отправлен", attachments=[self.get_back_keyboard()],
                                format=TextFormat.HTML)
        except Exception as e:
            await message.edit(text=f"❌ Ошибка: {str(e)[:50]}", attachments=[self.get_back_keyboard()],
                                format=TextFormat.HTML)

    async def show_instructions(self, message, user_id, context=None):
        text = (
            "<b>📝 ИНСТРУКЦИЯ ПО ДЕЖУРСТВУ</b>\n\n"
            "<b>▸ ПЕРЕД ДЕЖУРСТВОМ (пятница):</b>\n"
            "1. Позвонить в приемную: 5600 через вн. телефон в 17:00\n"
            "2. Сообщить о дежурстве и попросить оставить ключи на вахте\n\n"
            "<b>▸ В ДЕНЬ ДЕЖУРСТВА (суббота):</b>\n"
            "1. Прийти к 6:50 в АДЦ\n"
            "2. Взять ключ на охране от кубов\n"
            "3. Открыть кабинет 6002\n"
            "4. Сфотографировать открытый 6002 кабинет (как доказательство присутствия)\n"
            "5. Находиться там до 8:00\n"
            "6. После дежурства отписать в группу (пример: Доброе утро, никого из ЗГД не было)\n\n"
            "<b>▸ ОФОРМЛЕНИЕ ПРОТОКОЛА:</b>\n"
            "1. Распечатать бланк (предварительно написать дату)\n"
            "2. Расписаться на обороте:\n"
            "   ФИО, Должность, Модуль, Дата, Подпись\n"
            "3. Оставить у Е.С. Денисовой"
        )
        await message.edit(text=text, attachments=[self.get_back_keyboard()], format=TextFormat.HTML)

    async def show_questions(self, message, user_id, context=None):
        text = (
            "<b>❓ ЧАСТЫЕ ВОПРОСЫ</b>\n\n"
            "<b>▸ Не могу прийти на дежурство?</b>\n"
            "• Найти замену из списка\n"
            "• Сообщить М.С. Портновой\n"
            "• Пропуск = депремирование\n\n"
            "<b>▸ Ключ не на месте?</b>\n"
            "• Взять на охране ключ от теннисной переговорной\n"
            "• Сидеть возле кубов\n"
            "• В случае если пришёл ЗГД, провести в другую переговорную"
        )
        await message.edit(text=text, attachments=[self.get_back_keyboard()], format=TextFormat.HTML)

    async def change_profile(self, message, user_id, context=None):
        text = "<b>👤 ИЗМЕНЕНИЕ ПРОФИЛЯ</b>\n\nВыберите ваше ФИО из списка сотрудников.\n\n<i>Текущий выбор будет заменен.</i>"
        await message.edit(text=text, attachments=[self.get_employee_selection_keyboard(prefix="emp_")],
                            format=TextFormat.HTML)

    async def register_employee(self, message, user_id, employee_name: str):
        if user_id in self.user_data:
            self.user_data[user_id]["selected_employee"] = employee_name
            self.user_data[user_id]["registered_at"] = datetime.now().isoformat()
            self.save_user_data()

            text = (
                "<b>✅ РЕГИСТРАЦИЯ УСПЕШНА</b>\n\n"
                f"Ваш аккаунт привязан к:\n<b>{employee_name}</b>\n\n"
                f"📞 Телефон: {EMPLOYEE_PHONES.get(employee_name, 'не указан')}\n"
                f"🔔 Уведомления: {'✅ Включены' if self.user_data[user_id].get('notifications', True) else '❌ Отключены'}\n\n"
                "<i>Теперь вы можете пользоваться всеми функциями бота.</i>\n\nВыберите действие:"
            )
            await message.edit(text=text, attachments=[self.get_main_keyboard(user_id)], format=TextFormat.HTML)
        else:
            await message.edit(text="Ошибка регистрации. Пожалуйста, начните снова командой /start",
                                format=TextFormat.HTML)

    async def show_admin_panel(self, message, user_id, context=None):
        if not self.is_authorized_admin(user_id):
            await message.edit(text="❌ <b>ДОСТУП ЗАПРЕЩЕН</b>", format=TextFormat.HTML)
            return

        text = (
            "⚙️ <b>АДМИН-ПАНЕЛЬ</b>\n\n"
            "Доступные функции:\n\n"
            "📅 <b>Управление графиком:</b>\n"
            "• Добавить/удалить дежурство\n"
            "• Просмотреть график\n\n"
            "👥 <b>Управление сотрудниками:</b>\n"
            "• Добавить/удалить сотрудника\n"
            "• Изменить телефон\n"
            "• Список сотрудников\n\n"
            "📁 <b>Управление файлами:</b>\n"
            "• Загрузить протокол\n"
            "• Прикрепить протокол\n"
            "• Удалить файлы\n"
            "• Проверить файл\n\n"
            "📊 <b>Статистика:</b>\n"
            "• Активность пользователей\n"
            "• История действий\n\n"
            "<i>Выберите раздел:</i>"
        )
        await message.edit(text=text, attachments=[self.get_admin_keyboard()], format=TextFormat.HTML)

    async def admin_logout(self, message, user_id, context=None):
        if user_id in self.user_data:
            self.user_data[user_id]["is_admin"] = False
            self.user_data[user_id]["admin_logged_in"] = False
            self.save_user_data()

        await message.edit(
            text="✅ <b>ВЫ УСПЕШНО ВЫШЛИ ИЗ АДМИН-ПАНЕЛИ</b>\n\nВсе права администратора отозваны.",
            attachments=[self.get_main_keyboard(user_id)],
            format=TextFormat.HTML
        )

    async def admin_refresh_schedule(self, message, user_id, context=None):
        await self.show_admin_schedule(message, user_id, context)

    async def show_admin_schedule(self, message, user_id, context=None):
        text = (
            "📅 <b>УПРАВЛЕНИЕ ГРАФИКОМ ДЕЖУРСТВ</b>\n\n"
            "Доступные действия:\n\n"
            "➕ <b>Добавить дежурство:</b>\nСоздать новую ручную запись в графике поверх круга\n\n"
            "➖ <b>Удалить дежурство:</b>\nУдалить существующую ручную правку\n\n"
            "📋 <b>Просмотреть график:</b>\nПосмотреть текущий график\n\n"
            "🔄 <b>Обновить график:</b>\nОбновить отображение графика\n\n"
            "<i>Выберите действие:</i>"
        )
        await message.edit(text=text, attachments=[self.get_schedule_admin_keyboard()], format=TextFormat.HTML)

    async def show_admin_employees(self, message, user_id, context=None):
        text = (
            "👥 <b>УПРАВЛЕНИЕ СОТРУДНИКАМИ</b>\n\n"
            "Доступные действия:\n\n"
            "➕ <b>Добавить сотрудника:</b>\nДобавить нового сотрудника в систему\n\n"
            "➖ <b>Удалить сотрудника:</b>\nУдалить сотрудника из системы\n\n"
            "📞 <b>Изменить телефон:</b>\nОбновить контактный номер\n\n"
            "👥 <b>Список сотрудников:</b>\nПросмотреть всех сотрудников\n\n"
            "<i>Выберите действие:</i>"
        )
        await message.edit(text=text, attachments=[self.get_employees_admin_keyboard()], format=TextFormat.HTML)

    async def show_admin_files(self, message, user_id, context=None):
        protocol_exists = os.path.exists(self.protocol_file_path)
        text = (
            "📁 <b>УПРАВЛЕНИЕ ФАЙЛАМИ</b>\n\n"
            f"📄 <b>Протокол разногласий:</b>\n"
            f"Статус: {'✅ Доступен' if protocol_exists else '❌ Отсутствует'}\n"
            f"Закреплен: {'✅ Да' if self.protocol_pinned_message_id else '❌ Нет'}\n\n"
            "Доступные действия:\n\n"
            "📤 <b>Загрузить протокол:</b>\nДобавить новый файл протокола\n\n"
            "📎 <b>Прикрепить протокол:</b>\nЗакрепить сообщение с файлом в чате\n\n"
            "🗑 <b>Удалить протокол:</b>\nУдалить текущий файл протокола\n\n"
            "📄 <b>Проверить файл:</b>\nПроверить наличие и доступность\n\n"
            "<i>Выберите действие:</i>"
        )
        await message.edit(text=text, attachments=[self.get_files_admin_keyboard()], format=TextFormat.HTML)

    async def show_admin_stats(self, message, user_id, context=None):
        total_users = len(self.user_data)
        active_today = 0
        today = datetime.now().date()

        for user_info in self.user_data.values():
            last_active = user_info.get("last_active")
            if last_active:
                try:
                    last_active_date = datetime.fromisoformat(last_active).date()
                    if last_active_date == today:
                        active_today += 1
                except Exception:
                    pass

        auto_linked = 0
        for user_info in self.user_data.values():
            if user_info.get("selected_employee"):
                username = user_info.get("username", "")
                if username and self.get_employee_by_username(username):
                    auto_linked += 1

        next_saturday = None
        today_date = datetime.now(MOSCOW_TZ).replace(tzinfo=None).date()
        for date in range(1, 8):
            check_date = today_date + timedelta(days=date)
            if check_date.weekday() == 5:
                next_saturday = check_date
                break

        current_schedule = self.schedule_generator._generate_dynamic_schedule()
        next_duty = next((d for d in current_schedule.values() if d["date_obj"].date() == next_saturday), None)

        text = (
            "📊 <b>СТАТИСТИКА СИСТЕМЫ</b>\n\n"
            f"👥 <b>Всего пользователей:</b> {total_users}\n"
            f"📱 <b>Активных сегодня:</b> {active_today}\n"
            f"🤖 <b>Автопривязанных:</b> {auto_linked}\n"
            f"📅 <b>Дежурств на выводе:</b> {len(current_schedule)}\n"
            f"👥 <b>Ручных правок в базе:</b> {len(self.schedule_generator.schedule)}\n"
            f"👤 <b>Всего сотрудников:</b> {len(EMPLOYEE_PHONES)}\n\n"
        )

        if next_duty:
            text += f"<b>Следующее дежурство ({next_saturday.strftime('%d.%m.%Y')}):</b>\n"
            if next_duty["is_pair"]:
                text += f"• {next_duty['employees'][0]} + {next_duty['employees'][1]}\n"
            else:
                text += f"• {next_duty['employees'][0]}\n"
        else:
            text += f"<b>Ближайшая суббота ({next_saturday.strftime('%d.%m.%Y')}):</b>\n• Дежурных нет\n"

        text += f"\n<b>Расписание уведомлений (ВСЕМ):</b>\n• Среда 18:00 - уведомление о дежурстве в субботу\n• Пятница 18:00 - напоминание о завтрашнем дежурстве\n• Суббота 10:00 - напоминание в день дежурства\n"

        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="🔙 Назад в админку", payload="admin_panel"))
        kb.row(CallbackButton(text="🔄 Обновить", payload="admin_stats"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)

    async def admin_remove_duty(self, message, user_id, context: BaseContext):
        schedule_text = self.schedule_generator.get_schedule_text()
        text = (
                "➖ <b>УДАЛЕНИЕ РУЧНОГО ДЕЖУРСТВА</b>\n\n"
                "<i>Текущий график дежурств:</i>\n\n" +
                schedule_text[:1500] +
                "\n\nДля удаления ручной правки и возврата автоматического круга отправьте дату:\n"
                "<code>дд.мм.гггг г.</code>\n\n"
                "<b>Пример:</b> <code>06.06.2026г.</code>\n\n"
                "<i>Отправьте дату или нажмите 'Отмена':</i>"
        )
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="❌ Отмена", payload="admin_schedule"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
        await context.set_state(AdminWizard.awaiting_duty_remove)

    async def admin_add_employee(self, message, user_id, context: BaseContext):
        text = (
            "➕ <b>ДОБАВЛЕНИЕ СОТРУДНИКА</b>\n\n"
            "<i>Для добавления сотрудника отправьте сообщение в формате:</i>\n\n"
            "<code>ФИО;телефон;username</code>\n\n"
            "<b>Пример:</b>\n"
            "<code>Иванов Иван Иванович;8-999-111-11-11;@ivanov</code>\n\n"
            "<i>Важно:</i>\n"
            "• ФИО в формате: Фамилия И.О.\n"
            "• Телефон в формате: 8-XXX-XXX-XX-XX\n"
            "• Username в MAX с @ или без\n\n"
            "<i>Отправьте данные или нажмите 'Отмена':</i>"
        )
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="❌ Отмена", payload="admin_employees"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
        await context.set_state(AdminWizard.awaiting_employee_add)

    async def admin_remove_employee(self, message, user_id, context: BaseContext):
        employees_list = "\n".join([f"• {emp}" for emp in EMPLOYEE_PHONES.keys()])
        text = (
            "➖ <b>УДАЛЕНИЕ СОТРУДНИКА</b>\n\n"
            f"<b>Список сотрудников:</b>\n{employees_list}\n\n"
            "<i>Для удаления сотрудника отправьте его ФИО:</i>\n\n"
            "<b>Пример:</b>\n<code>Иванов И.И.</code>\n\n"
            "<i>Отправьте ФИО или нажмите 'Отмена':</i>"
        )
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="❌ Отмена", payload="admin_employees"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
        await context.set_state(AdminWizard.awaiting_employee_remove)

    async def admin_edit_phone(self, message, user_id, context: BaseContext):
        employees_list = "\n".join([f"• {emp}" for emp in EMPLOYEE_PHONES.keys()])
        text = (
            "📞 <b>ИЗМЕНЕНИЕ ТЕЛЕФОНА СОТРУДНИКА</b>\n\n"
            f"<b>Список сотрудников:</b>\n{employees_list}\n\n"
            "<i>Для изменения телефона отправьте сообщение в формате:</i>\n\n"
            "<code>ФИО;новый телефон</code>\n\n"
            "<b>Пример:</b>\n<code>Иванов И.И.;8-900-000-00-00</code>\n\n"
            "<i>Отправьте данные или нажмите 'Отмена':</i>"
        )
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="❌ Отмена", payload="admin_employees"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
        await context.set_state(AdminWizard.awaiting_phone_edit)

    async def admin_list_employees(self, message, user_id, context=None):
        employees_text = ""
        for i, (employee, phone) in enumerate(EMPLOYEE_PHONES.items(), 1):
            username = None
            for uname, emp_name in USERNAME_TO_EMPLOYEE.items():
                if emp_name == employee:
                    username = uname
                    break
            employees_text += f"{i}. <b>{employee}</b>\n   📞 {phone}\n"
            if username:
                employees_text += f"   📱 MAX: {username}\n"
            employees_text += "\n"

        text = f"👥 <b>СПИСОК СОТРУДНИКОВ</b>\n\n{employees_text}<b>Всего сотрудников:</b> {len(EMPLOYEE_PHONES)}"
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="➕ Добавить сотрудника", payload="admin_add_employee"))
        kb.row(CallbackButton(text="📞 Изменить телефон", payload="admin_edit_phone"))
        kb.row(CallbackButton(text="🔙 Назад", payload="admin_employees"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)

    async def admin_upload_protocol(self, message, user_id, context: BaseContext):
        text = (
            "📤 <b>ЗАГРУЗКА ПРОТОКОЛА</b>\n\n"
            "Пришлите следующим сообщением файл протокола (без подписи — "
            "в MAX нельзя прикрепить текст к файлу).\n\n"
            "Бот сохранит именно следующий присланный файл.\n\n"
            "<b>Формат файла:</b> .docx\n"
            "<b>Рекомендуемое имя:</b> Протокол разногласий — пример.docx"
        )
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="❌ Отмена", payload="admin_files"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
        await context.set_state(AdminWizard.awaiting_protocol_upload)

    async def admin_delete_protocol(self, message, user_id, context=None):
        if os.path.exists(self.protocol_file_path):
            try:
                os.remove(self.protocol_file_path)
                self.protocol_pinned_message_id = None
                text = "🗑 <b>ФАЙЛ ПРОТОКОЛА УДАЛЕН</b>\n\nФайл протокола был успешно удален.\n\n<i>Пользователи больше не смогут скачать протокол.</i>"
            except Exception as e:
                text = f"❌ <b>ОШИБКА УДАЛЕНИЯ:</b> {str(e)}"
        else:
            text = "ℹ️ <b>ФАЙЛ НЕ НАЙДЕН</b>\n\nФайл протокола уже отсутствует."

        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="🔙 Назад", payload="admin_files"))
        kb.row(CallbackButton(text="📄 Проверить файл", payload="admin_check_protocol"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)

    async def admin_check_protocol(self, message, user_id, context=None):
        protocol_exists = os.path.exists(self.protocol_file_path)
        if protocol_exists:
            file_size = os.path.getsize(self.protocol_file_path)
            file_size_mb = file_size / (1024 * 1024)
            text = (
                "✅ <b>ФАЙЛ ПРОТОКОЛА НАЙДЕН</b>\n\n"
                f"📄 <b>Имя файла:</b> {os.path.basename(self.protocol_file_path)}\n"
                f"📁 <b>Размер:</b> {file_size_mb:.2f} МБ\n"
                f"📍 <b>Путь:</b> {self.protocol_file_path}\n"
                f"📎 <b>Закреплен:</b> {'Да' if self.protocol_pinned_message_id else 'Нет'}\n\n"
                "<i>Файл доступен для скачивания пользователями.</i>"
            )
        else:
            text = (
                "❌ <b>ФАЙЛ ПРОТОКОЛА НЕ НАЙДЕН</b>\n\n"
                f"<i>Путь:</i> {self.protocol_file_path}\n\n"
                "<b>Что делать:</b>\n1. Загрузите файл протокола\n2. Используйте кнопку 'Загрузить протокол'"
            )

        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="🔙 Назад", payload="admin_files"))
        kb.row(CallbackButton(text="📤 Загрузить протокол", payload="admin_upload_protocol"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)

    async def admin_pin_protocol(self, message, user_id, context: BaseContext):
        if not os.path.exists(self.protocol_file_path):
            text = "❌ <b>ФАЙЛ НЕ НАЙДЕН</b>\n\nСначала загрузите файл протокола.\nИспользуйте кнопку 'Загрузить протокол'."
            kb = InlineKeyboardBuilder()
            kb.row(CallbackButton(text="📤 Загрузить протокол", payload="admin_upload_protocol"))
            kb.row(CallbackButton(text="🔙 Назад", payload="admin_files"))
            await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
            return

        text = (
            "📎 <b>ПРИКРЕПЛЕНИЕ ПРОТОКОЛА</b>\n\n"
            "Пришлите следующим сообщением файл протокола (без подписи) — "
            "бот закрепит именно это сообщение в чате."
        )
        kb = InlineKeyboardBuilder()
        kb.row(CallbackButton(text="❌ Отмена", payload="admin_files"))
        await message.edit(text=text, attachments=[kb.as_markup()], format=TextFormat.HTML)
        await context.set_state(AdminWizard.awaiting_protocol_pin)

    # ================= ПОШАГОВЫЙ МАСТЕР (текстовые шаги) =================

    async def wizard_wait_date(self, event: MessageCreated, context: BaseContext):
        message_text = event.message.body.text or ""
        data = await context.get_data()
        duty = data.get("new_duty", {"is_pair": False, "employees": [], "phones": []})
        duty["date"] = message_text if "г." in message_text else message_text + "г."
        await context.update_data(new_duty=duty)
        await context.set_state(None)

        await event.message.answer(
            f"✅ Дата сохранена: <b>{message_text}</b>\n\n👤 <b>Шаг 3:</b> Выберите первого дежурного из списка:",
            attachments=[self.get_employee_selection_keyboard("add_e1_")],
            format=TextFormat.HTML
        )

    async def wizard_wait_phones(self, event: MessageCreated, context: BaseContext):
        message_text = event.message.body.text or ""
        data = await context.get_data()
        duty_info = data["new_duty"]

        if message_text.lower() in ["ок", "ok", "да", "авто"]:
            phones = [EMPLOYEE_PHONES.get(emp, "Номер не найден") for emp in duty_info['employees']]
        else:
            phones = [p.strip() for p in message_text.split(',')]

        if len(duty_info['employees']) != len(phones):
            await event.message.answer(
                "❌ <b>ОШИБКА</b>\n\nКоличество сотрудников и телефонов не совпадает. Введите телефоны заново:",
                format=TextFormat.HTML
            )
            return

        success, msg = self.schedule_generator.add_duty(
            duty_info['date'], duty_info['employees'], phones, duty_info['is_pair']
        )

        if success:
            audit(str(event.message.sender.user_id), event.message.sender.username, "duty_added",
                  f"{duty_info['date']}: {', '.join(duty_info['employees'])}")
            await event.message.answer(
                f"✅ <b>РУЧНОЕ ДЕЖУРСТВО ДОБАВЛЕНО</b>\n\n"
                f"📅 Дата: {duty_info['date']}\n"
                f"👥 Сотрудники: {', '.join(duty_info['employees'])}\n"
                f"📞 Телефоны: {', '.join(phones)}\n\n"
                "<i>График успешно обновлен (круг перезаписан на этот день).</i>",
                attachments=[self.get_admin_keyboard()],
                format=TextFormat.HTML
            )
        else:
            await event.message.answer(
                f"❌ <b>ОШИБКА ДОБАВЛЕНИЯ</b>\n\n{msg}",
                attachments=[self.get_admin_keyboard()],
                format=TextFormat.HTML
            )

        await context.clear()

    async def wizard_duty_remove(self, event: MessageCreated, context: BaseContext):
        date_str = (event.message.body.text or "").strip()
        success = self.schedule_generator.remove_duty(date_str)

        if success:
            audit(str(event.message.sender.user_id), event.message.sender.username, "duty_removed", date_str)
            await event.message.answer(
                f"✅ <b>РУЧНАЯ ПРАВКА УДАЛЕНА</b>\n\n📅 Дата: {date_str}\n\n<i>На этот день вернулся автоматический расчет по кругу.</i>",
                format=TextFormat.HTML
            )
        else:
            await event.message.answer(
                f"❌ <b>РУЧНАЯ ПРАВКА НЕ НАЙДЕНА</b>\n\nДата: {date_str}\n\nПроверьте правильность введённой даты.",
                format=TextFormat.HTML
            )
        await context.clear()

    async def wizard_employee_add(self, event: MessageCreated, context: BaseContext):
        message_text = event.message.body.text or ""
        try:
            parts = message_text.split(';')
            if len(parts) == 3:
                employee_name = parts[0].strip()
                phone = parts[1].strip()
                username = parts[2].strip()

                success = self.schedule_generator.add_employee(employee_name, phone)

                if success:
                    if username:
                        if not username.startswith('@'):
                            username = '@' + username
                        USERNAME_TO_EMPLOYEE[username.lower()] = employee_name

                    audit(str(event.message.sender.user_id), event.message.sender.username,
                          "employee_added", employee_name)
                    await event.message.answer(
                        f"✅ <b>СОТРУДНИК ДОБАВЛЕН</b>\n\n👤 ФИО: {employee_name}\n📞 Телефон: {phone}\n📱 MAX: {username if username else 'не указан'}\n\n<i>Сотрудник добавлен в систему.</i>",
                        format=TextFormat.HTML
                    )
                else:
                    await event.message.answer(
                        f"❌ <b>СОТРУДНИК УЖЕ СУЩЕСТВУЕТ</b>\n\nИмя: {employee_name}\n\nИспользуйте другое ФИО.",
                        format=TextFormat.HTML
                    )
            else:
                await event.message.answer(
                    "❌ <b>НЕВЕРНЫЙ ФОРМАТ</b>\n\nИспользуйте формат:\n<code>ФИО;телефон;username</code>",
                    format=TextFormat.HTML)
        except Exception as e:
            await event.message.answer(f"❌ <b>ОШИБКА:</b> {str(e)}\n\nПроверьте правильность данных.",
                                        format=TextFormat.HTML)
        await context.clear()

    async def wizard_employee_remove(self, event: MessageCreated, context: BaseContext):
        employee_name = (event.message.body.text or "").strip()
        success = self.schedule_generator.remove_employee(employee_name)

        if success:
            usernames = []
            for uname, emp_name in list(USERNAME_TO_EMPLOYEE.items()):
                if emp_name == employee_name:
                    usernames.append(uname)
                    del USERNAME_TO_EMPLOYEE[uname]

            username_info = f"\n📱 MAX: {', '.join(usernames)}" if usernames else ""

            audit(str(event.message.sender.user_id), event.message.sender.username,
                  "employee_removed", employee_name)
            await event.message.answer(
                f"✅ <b>СОТРУДНИК УДАЛЕН</b>\n\n👤 ФИО: {employee_name}{username_info}\n\n<i>Сотрудник удален из системы.</i>",
                format=TextFormat.HTML
            )
        else:
            await event.message.answer(
                f"❌ <b>СОТРУДНИК НЕ НАЙДЕН</b>\n\nИмя: {employee_name}\n\nПроверьте правильность ФИО.",
                format=TextFormat.HTML
            )
        await context.clear()

    async def wizard_phone_edit(self, event: MessageCreated, context: BaseContext):
        message_text = event.message.body.text or ""
        try:
            parts = message_text.split(';')
            if len(parts) == 2:
                employee_name = parts[0].strip()
                new_phone = parts[1].strip()

                success = self.schedule_generator.update_employee_phone(employee_name, new_phone)

                if success:
                    audit(str(event.message.sender.user_id), event.message.sender.username,
                          "phone_updated", f"{employee_name} -> {new_phone}")
                    await event.message.answer(
                        f"✅ <b>ТЕЛЕФОН ОБНОВЛЕН</b>\n\n👤 Сотрудник: {employee_name}\n📞 Новый телефон: {new_phone}\n\n<i>Телефон успешно обновлен.</i>",
                        format=TextFormat.HTML
                    )
                else:
                    await event.message.answer(
                        f"❌ <b>СОТРУДНИК НЕ НАЙДЕН</b>\n\nИмя: {employee_name}\n\nПроверьте правильность ФИО.",
                        format=TextFormat.HTML)
            else:
                await event.message.answer(
                    "❌ <b>НЕВЕРНЫЙ ФОРМАТ</b>\n\nИспользуйте формат:\n<code>ФИО;новый телефон</code>",
                    format=TextFormat.HTML)
        except Exception as e:
            await event.message.answer(f"❌ <b>ОШИБКА:</b> {str(e)}\n\nПроверьте правильность данных.",
                                        format=TextFormat.HTML)
        await context.clear()

    # ================= ФАЙЛЫ ПРОТОКОЛА (входящие вложения) =================

    @staticmethod
    def _get_attachment_filename(attachment) -> str:
        return getattr(attachment, "filename", None) or ""

    @staticmethod
    def _get_attachment_url(attachment) -> Optional[str]:
        payload = getattr(attachment, "payload", None)
        return getattr(payload, "url", None)

    async def wizard_protocol_upload(self, event: MessageCreated, context: BaseContext):
        """Файл, присланный сразу после кнопки «Загрузить протокол» (без подписи —
        в MAX нельзя прикрепить текст к файлу, поэтому используем состояние мастера)."""
        body = event.message.body
        attachments = body.attachments if body else None
        await context.clear()

        if not attachments:
            await event.message.answer(
                "❌ <b>ЭТО НЕ ФАЙЛ</b>\n\nНажмите «Загрузить протокол» ещё раз и пришлите файл .docx.",
                format=TextFormat.HTML
            )
            return

        attachment = attachments[0]
        filename = self._get_attachment_filename(attachment)

        if not filename.endswith(".docx"):
            await event.message.answer(
                "❌ <b>НЕВЕРНЫЙ ФОРМАТ ФАЙЛА</b>\n\nПоддерживаются только файлы .docx",
                format=TextFormat.HTML
            )
            return

        url = self._get_attachment_url(attachment)
        if not url:
            await event.message.answer("❌ <b>ОШИБКА ЗАГРУЗКИ:</b> файл недоступен для скачивания",
                                        format=TextFormat.HTML)
            return

        try:
            downloaded_path = await self.bot.download_file(url, self.download_dir)
            shutil.move(str(downloaded_path), self.protocol_file_path)
            size_kb = os.path.getsize(self.protocol_file_path) / 1024
            audit(str(event.message.sender.user_id), event.message.sender.username,
                  "protocol_uploaded", filename)
            await event.message.answer(
                f"✅ <b>ФАЙЛ ПРОТОКОЛА ЗАГРУЖЕН</b>\n\n📄 Имя файла: {filename}\n📁 Размер: {size_kb:.1f} КБ\n\n<i>Файл успешно сохранен и доступен для скачивания.</i>",
                format=TextFormat.HTML
            )
        except Exception as e:
            await event.message.answer(f"❌ <b>ОШИБКА ЗАГРУЗКИ:</b> {str(e)}", format=TextFormat.HTML)

    async def wizard_protocol_pin(self, event: MessageCreated, context: BaseContext):
        """Файл, присланный сразу после кнопки «Прикрепить протокол»."""
        body = event.message.body
        attachments = body.attachments if body else None
        await context.clear()

        if not attachments:
            await event.message.answer(
                "❌ <b>ЭТО НЕ ФАЙЛ</b>\n\nНажмите «Прикрепить протокол» ещё раз и пришлите файл .docx.",
                format=TextFormat.HTML
            )
            return

        filename = self._get_attachment_filename(attachments[0])
        if not filename.endswith(".docx"):
            await event.message.answer(
                "❌ <b>НЕВЕРНЫЙ ФОРМАТ ФАЙЛА</b>\n\nПоддерживаются только файлы .docx",
                format=TextFormat.HTML
            )
            return

        try:
            chat_id = event.message.recipient.chat_id
            mid = body.mid
            await self.bot.pin_message(chat_id=chat_id, message_id=mid)
            self.protocol_pinned_message_id = mid
            audit(str(event.message.sender.user_id), event.message.sender.username,
                  "protocol_pinned", filename)
            await event.message.answer(
                f"✅ <b>ФАЙЛ ПРОТОКОЛА ПРИКРЕПЛЕН</b>\n\n📄 Имя файла: {filename}\n\n<i>Сообщение с файлом закреплено в чате.</i>",
                format=TextFormat.HTML
            )
        except Exception as e:
            await event.message.answer(f"❌ <b>ОШИБКА ПРИКРЕПЛЕНИЯ:</b> {str(e)}", format=TextFormat.HTML)

    async def on_document(self, event: MessageCreated):
        """Файл прислан вне мастера (например, случайно) — просто подсказываем, как загрузить протокол."""
        sender = event.message.sender
        user_id = str(sender.user_id) if sender else None
        if not user_id or not self.is_admin(user_id):
            return  # обычный пользователь прислал файл — молча игнорируем

        if not self.is_authorized_admin(user_id):
            await event.message.answer(
                "❌ <b>СЕССИЯ АДМИНИСТРАТОРА НЕ АКТИВНА</b>\n\n"
                "Бот мог перезапуститься на хостинге — войдите заново:\n"
                "<code>/admin логин пароль</code>",
                format=TextFormat.HTML
            )
            return

        await event.message.answer(
            "ℹ️ <b>ФАЙЛ ПОЛУЧЕН, НО НЕ ОБРАБОТАН</b>\n\n"
            "Чтобы бот сохранил или закрепил файл — сначала откройте "
            "<i>Админ-панель → Управление файлами</i> и нажмите нужную кнопку, "
            "а уже потом присылайте файл.",
            format=TextFormat.HTML
        )

    # ================= ФОЛБЭК ДЛЯ ОБЫЧНОГО ТЕКСТА =================

    async def on_plain_text(self, event: MessageCreated):
        body = event.message.body
        message_text = (body.text if body else "") or ""

        if message_text.startswith('/'):
            return

        sender = event.message.sender
        user_id = str(sender.user_id) if sender else None

        if user_id and self.is_authorized_admin(user_id):
            await event.message.answer(
                "ℹ️ <b>ИНФОРМАЦИЯ</b>\n\nДля работы с админ-панелью используйте кнопки меню.\nИли вернитесь в главное меню.",
                attachments=[self.get_admin_keyboard()],
                format=TextFormat.HTML
            )
        else:
            await event.message.answer(
                "ℹ️ <b>ИНФОРМАЦИЯ</b>\n\n"
                "Я бот для управления графиком дежурств.\n"
                "Используйте кнопки меню для навигации.\n\n"
                "<i>Для админ-функций необходимо войти в админ-панель.</i>",
                format=TextFormat.HTML
            )

    # ================= ЗАПУСК =================

    async def _run_async(self):
        await self._start_scheduler()
        logger.info("Бот запущен...")
        await self.dp.start_polling(self.bot)

    def run(self):
        asyncio.run(self._run_async())


if __name__ == "__main__":
    bot = DutyBot()
    bot.run()
